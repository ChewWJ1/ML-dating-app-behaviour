import sys
import re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx')

print("=== SEARCH FOR '16 models' or '14 models' ===")
for i, p in enumerate(doc.paragraphs):
    txt = p.text
    if re.search(r'\b1[46]\b.*\bmodels?\b', txt, re.IGNORECASE):
        print(f"Para {i}: {txt}")
        print("-" * 50)

for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            txt = cell.text
            if re.search(r'\b1[46]\b.*\bmodels?\b', txt, re.IGNORECASE):
                print(f"Table {t_idx}, Row {r_idx}, Col {c_idx}: {txt}")
                print("-" * 50)
