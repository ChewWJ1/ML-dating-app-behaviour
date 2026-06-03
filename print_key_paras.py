import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

DOCX_PATH = r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx'
doc = Document(DOCX_PATH)

# Print FULL text of all key paragraphs identified in previous audit
para_ids = [27, 28, 57, 67, 103, 105, 120, 121, 122, 123, 124, 125, 126, 127, 128, 
            129, 130, 131, 132, 133, 140, 143, 144, 145, 150, 151, 152, 153, 
            259, 291, 346, 353, 361, 373, 395]

for pi in para_ids:
    if pi < len(doc.paragraphs):
        t = doc.paragraphs[pi].text
        if t.strip():
            print(f"Para {pi}: {t}")
            print("-" * 60)
