"""
Targeted audit script for 5 specific discrepancies reported by user:
1. Feature counts (116 vs 122 input features; 67 vs 66 selected)
2. Engineered feature names (fabricated vs real)
3. DML p-value contradiction (p > 0.60 vs p = 0.0322)
4. Gender parity accuracy values not in notebook
5. PCA component count (55 vs 24)
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

DOCX_PATH = r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx'
doc = Document(DOCX_PATH)

print("=" * 70)
print("TARGETED AUDIT FOR 5 DISCREPANCIES")
print("=" * 70)

# ───────────────────────────────────────────────────────────
print("\n--- DISCREPANCY 1: Feature counts (116 vs 122 input, 67 vs 66 selected) ---")
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if any(x in t for x in ["116 input", "116 features", "67 features", "67 selected",
                              "122 features", "122 input", "66 features", "66 selected",
                              "expanding features from 25", "selected features", "feature selection"]):
        print(f"  Para {i}: {t}")
        print()

for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        seen = set()
        for ci, cell in enumerate(row.cells):
            if cell in seen: continue
            seen.add(cell)
            t = cell.text
            if any(x in t for x in ["116 input", "116 features", "67 features", "67 selected",
                                      "122 features", "66 features", "66 selected",
                                      "expanding features from 25", "feature selection"]):
                print(f"  Table {ti}, R{ri}, C{ci}: {t}")
                print()

# ───────────────────────────────────────────────────────────
print("\n--- DISCREPANCY 2: Engineered feature names ---")
keywords = ["popularity_density", "bio_message_interaction", "selective_emoji_swiper",
            "engagement_score", "profile_completeness", "activity_intensity",
            "selectivity_ratio", "late_night_user", "interaction feature"]
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if any(k in t.lower() for k in [kw.lower() for kw in keywords]):
        print(f"  Para {i}: {t}")
        print()

for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        seen = set()
        for ci, cell in enumerate(row.cells):
            if cell in seen: continue
            seen.add(cell)
            t = cell.text
            if any(k in t.lower() for k in [kw.lower() for kw in keywords]):
                print(f"  Table {ti}, R{ri}, C{ci}: {t}")
                print()

# ───────────────────────────────────────────────────────────
print("\n--- DISCREPANCY 3: DML p-value contradiction ---")
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if any(x in t for x in ["p > 0.60", "p>0.60", "0.0322", "indistinguishable from zero",
                              "ATE is", "treatment effect", "p = 0", "p-value"]):
        print(f"  Para {i}: {t}")
        print()

for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        seen = set()
        for ci, cell in enumerate(row.cells):
            if cell in seen: continue
            seen.add(cell)
            t = cell.text
            if any(x in t for x in ["p > 0.60", "p>0.60", "0.0322", "indistinguishable from zero",
                                      "ATE is", "treatment effect", "p = 0", "p-value"]):
                print(f"  Table {ti}, R{ri}, C{ci}: {t}")
                print()

# ───────────────────────────────────────────────────────────
print("\n--- DISCREPANCY 4: Gender parity accuracy values ---")
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if any(x in t for x in ["59.4", "60.2", "60.1", "Male (", "Female (", "Non-binary",
                              "gender", "parity", "fairness", "Transgender", "TPR", "FPR"]):
        print(f"  Para {i}: {t}")
        print()

for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        seen = set()
        for ci, cell in enumerate(row.cells):
            if cell in seen: continue
            seen.add(cell)
            t = cell.text
            if any(x in t for x in ["59.4", "60.2", "60.1", "Male (", "Female (", "Non-binary",
                                      "gender", "parity", "fairness", "Transgender"]):
                print(f"  Table {ti}, R{ri}, C{ci}: {t}")
                print()

# ───────────────────────────────────────────────────────────
print("\n--- DISCREPANCY 5: PCA component count (55 vs 24) ---")
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if any(x in t for x in ["55 principal", "55 component", "PCA", "principal component",
                              "24 component", "24 principal", "95% variance"]):
        print(f"  Para {i}: {t}")
        print()

for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        seen = set()
        for ci, cell in enumerate(row.cells):
            if cell in seen: continue
            seen.add(cell)
            t = cell.text
            if any(x in t for x in ["55 principal", "55 component", "PCA", "principal component",
                                      "24 component", "24 principal", "95% variance"]):
                print(f"  Table {ti}, R{ri}, C{ci}: {t}")
                print()

print("\n" + "=" * 70)
print("AUDIT COMPLETE")
print("=" * 70)
