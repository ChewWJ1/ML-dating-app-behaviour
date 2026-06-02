import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

root_dir = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour"
doc_path = os.path.join(root_dir, "reports", "WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx")

if not os.path.exists(doc_path):
    print("Error: Report file not found!")
    sys.exit(1)

doc = docx.Document(doc_path)

print("Checking paragraphs...")
paragraph_matches = []
for idx, p in enumerate(doc.paragraphs):
    text = p.text
    # Search for terms
    for term in ["XGBoost", "Random Forest", "champion", "Brier"]:
        if term.lower() in text.lower():
            paragraph_matches.append((idx, term, text))

print(f"Found {len(paragraph_matches)} paragraph matches:")
for idx, term, text in paragraph_matches:
    print(f"  P {idx} (term: {term}): {text[:150]}...")

print("\nChecking tables...")
table_matches = []
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text
            for term in ["XGBoost", "Random Forest", "champion", "Brier"]:
                if term.lower() in text.lower():
                    table_matches.append((t_idx, r_idx, c_idx, term, text))

print(f"Found {len(table_matches)} table matches:")
for t_idx, r_idx, c_idx, term, text in table_matches:
    # Clean text display
    clean_text = " | ".join([line.strip() for line in text.split("\n") if line.strip()])
    print(f"  Table {t_idx}, Row {r_idx}, Cell {c_idx} (term: {term}): {clean_text[:120]}...")
