import docx
import os

doc_path = "reports/WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5 SOTA.docx"
if not os.path.exists(doc_path):
    print("Document not found!")
    exit(1)

doc = docx.Document(doc_path)
print(f"Document: {doc_path}")
print(f"Total Paragraphs: {len(doc.paragraphs)}")
print(f"Total Tables: {len(doc.tables)}")

word_count = sum(len(p.text.split()) for p in doc.paragraphs)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            word_count += len(cell.text.split())
print(f"Approximate Word Count: {word_count}")

print("\n--- Document Headings ---")
for idx, p in enumerate(doc.paragraphs):
    if p.style.name.startswith("Heading") or (p.text.strip() and p.text.strip()[0].isdigit() and p.text.split()[0].replace('.', '').isdigit()):
        print(f"P {idx} ({p.style.name}): {p.text[:100]}")
