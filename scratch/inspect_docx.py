import docx

doc_path = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report.docx"
doc = docx.Document(doc_path)

print("Total Paragraphs:", len(doc.paragraphs))
print("Total Tables:", len(doc.tables))

for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        # Print headings (which usually have bold or specific styles or just print any non-empty paragraph under 120 chars)
        if len(p.text) < 120:
            print(f"[{i}] Style: {p.style.name} | Text: {p.text}")
