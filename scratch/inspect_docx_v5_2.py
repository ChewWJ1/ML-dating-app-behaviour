import docx
import os

doc_path = r"reports/WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.2 SOTA.docx"
if not os.path.exists(doc_path):
    print("Error: File not found!")
    exit(1)

doc = docx.Document(doc_path)
print(f"Total paragraphs in V5.2 SOTA: {len(doc.paragraphs)}")

# Print paragraphs that look like headings or contain "Streamlit" or "SwipeIQ"
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text.startswith(("1.0", "2.0", "3.0", "4.0", "5.0", "6.0", "7.0", "8.0", "9.0")) or "Streamlit" in text or "SwipeIQ" in text:
        print(f"P {idx} ({p.style.name}): {text}")
