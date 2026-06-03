"""
Final definitive figure ordering fix.
Target order for each figure (matching rest of document convention):
  intro_sentence
  [IMG paragraph]
  Figure N: Title caption
  Description body text
  [blank]

Move each block precisely using this exact order.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from lxml import etree

DOCX_PATH = r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx'

def has_image(p):
    xml = etree.tostring(p._element, encoding='unicode')
    return '<w:drawing' in xml or '<v:shape' in xml

def move_ordered(doc, ordered_para_texts_or_indices, anchor_text, use_indices=False):
    """
    For each item in ordered_para_texts_or_indices (list of strings to match or list of indices),
    find the paragraph, remove it, then reinsert all in order before anchor_text paragraph.
    """
    paras = doc.paragraphs
    # Resolve to elements
    elems = []
    for item in ordered_para_texts_or_indices:
        if use_indices:
            elems.append(paras[item]._element)
        else:
            found = None
            for p in paras:
                if item == 'IMG':
                    # This is handled separately
                    break
                if item in p.text:
                    found = p._element
                    break
            if found is not None:
                elems.append(found)

    # Find anchor
    anchor_elem = None
    for p in paras:
        if anchor_text in p.text:
            anchor_elem = p._element
            break
    if not anchor_elem:
        print(f"  Anchor not found: {anchor_text[:50]}")
        return

    # Remove all
    for e in elems:
        e.getparent().remove(e)

    # Reinsert in correct order (reversed so addprevious works correctly)
    for e in reversed(elems):
        anchor_elem.addprevious(e)
    print(f"  Done: {len(elems)} elements placed before '{anchor_text[:50]}'")


def fix_block(doc, anchor_text, intro_text, caption_text, body_text, img_search_near_caption=True):
    """
    Find the 4 elements of a figure block and reinsert in correct order:
    intro -> [IMG] -> caption -> body
    """
    paras = doc.paragraphs
    intro_elem = caption_elem = body_elem = img_elem = None

    for p in paras:
        t = p.text
        if intro_text and intro_text in t:      intro_elem   = p._element
        elif caption_text and caption_text in t: caption_elem = p._element
        elif body_text and body_text in t:       body_elem    = p._element

    # Find image: blank paragraph with drawing, closest to caption
    caption_idx = next((i for i, p in enumerate(paras) if caption_text in p.text), None)
    if caption_idx is not None:
        candidates = [(abs(i - caption_idx), i) for i, p in enumerate(paras)
                      if has_image(p) and p.text.strip() == '' and abs(i - caption_idx) <= 6]
        if candidates:
            img_idx = sorted(candidates)[0][1]
            img_elem = paras[img_idx]._element

    anchor_elem = None
    for p in paras:
        if anchor_text in p.text:
            anchor_elem = p._element
            break
    if not anchor_elem:
        print(f"  Anchor not found: {anchor_text}")
        return

    # Collect in correct order: intro, img, caption, body
    elems = [e for e in [intro_elem, img_elem, caption_elem, body_elem] if e is not None]
    for e in elems:
        e.getparent().remove(e)
    for e in reversed(elems):
        anchor_elem.addprevious(e)

    found = [x is not None for x in [intro_elem, img_elem, caption_elem, body_elem]]
    print(f"  Placed [intro={found[0]}, img={found[1]}, cap={found[2]}, body={found[3]}] before '{anchor_text[:50]}'")


# ══════════════════════════════════════════════════════════════════════════════
doc = Document(DOCX_PATH)

print("=== FIG 41 ===")
fix_block(doc,
    anchor_text   = "4.3 Hyperparameter Tuning",
    intro_text    = "Figure 41 below presents the training loss comparison",
    caption_text  = "Figure 41: PyTorch BCE Loss Comparison",
    body_text     = "Figure 41 compares the training loss curves of standard Binary"
)
doc.save(DOCX_PATH); doc = Document(DOCX_PATH)

print("\n=== FIG 42 ===")
fix_block(doc,
    anchor_text   = "5.0 Insights and Interpretation",
    intro_text    = "Figure 42 below tracks the Opacus DP-SGD privacy budget",
    caption_text  = "Figure 42: Opacus Differential Privacy",
    body_text     = "Figure 42 tracks the privacy budget consumption (epsilon)"
)
doc.save(DOCX_PATH); doc = Document(DOCX_PATH)

print("\n=== FIG 39 ===")
fix_block(doc,
    anchor_text   = "5.3 Demographic Parity",
    intro_text    = "Complementing the SHAP attributions above",   # <-- intro goes first
    caption_text  = "Figure 39: Attentive TabNet",
    body_text     = "attentive feature selection heatmap in Figure 39"
)
doc.save(DOCX_PATH); doc = Document(DOCX_PATH)

print("\n=== FIG 40 (after Fig 39, before 5.3) ===")
fix_block(doc,
    anchor_text   = "5.3 Demographic Parity",
    intro_text    = None,  # no second intro
    caption_text  = "Figure 40: SCARF Self-Supervised",
    body_text     = "Figure 40 projects the SCARF contrastive pre-trained"
)
# Also move the SCARF disclosure just after Fig 40 body
paras = doc.paragraphs
disc_elem = anchor_elem = None
for p in paras:
    if "Methodological Disclosure: During our SCARF" in p.text: disc_elem = p._element
    if "5.3 Demographic Parity" in p.text:                      anchor_elem = p._element
if disc_elem and anchor_elem:
    disc_elem.getparent().remove(disc_elem)
    anchor_elem.addprevious(disc_elem)
    print("  SCARF disclosure moved before 5.3")
doc.save(DOCX_PATH); doc = Document(DOCX_PATH)

print("\n=== FIG 38 ===")
fix_block(doc,
    anchor_text   = "7.3 Summary of Evaluated and Excluded",
    intro_text    = "Figure 38 below presents the training loss curves comparing",
    caption_text  = "Figure 38: Knowledge Distillation Student",
    body_text     = "Figure 38 compares the training loss profiles of the teacher ensemble"
)
doc.save(DOCX_PATH)

# ══════════════════════════════════════════════════════════════════════════════
print("\n=== FINAL VERIFICATION ===")
doc = Document(DOCX_PATH); paras = doc.paragraphs

def show(label, keywords):
    print(f"\n{label}")
    for i, p in enumerate(paras):
        if any(kw in p.text for kw in keywords):
            img = '[IMG]' if has_image(p) else '     '
            print(f"  {i:4d} {img} {p.text[:105]}")

show("Fig 41 (should be: intro -> [IMG] -> caption -> body -> 4.3):",
     ["Figure 41 below", "Figure 41: PyTorch", "Figure 41 compares", "4.3 Hyperparameter"])

show("Fig 42 (should be: intro -> [IMG] -> caption -> body -> 5.0):",
     ["Figure 42 below", "Figure 42: Opacus", "Figure 42 tracks", "5.0 Insights"])

show("Figs 39+40 (intro -> cap39 -> [IMG]39 -> body39 -> cap40 -> [IMG]40 -> body40 -> disclosure -> 5.3):",
     ["Complementing the SHAP", "Figure 39: Attentive", "attentive feature selection", 
      "Figure 40: SCARF", "Figure 40 projects", "Methodological Disclosure: During our SCARF", "5.3 Demographic"])

show("Fig 38 (intro -> [IMG] -> caption -> body -> 7.3):",
     ["Figure 38 below", "Figure 38: Knowledge", "Figure 38 compares", "7.3 Summary of Evaluated"])

print("\n-- Section 8 contents (should have zero figure 38-42 paragraphs) --")
in8 = False
for i, p in enumerate(paras):
    if "8.1 Key Findings" in p.text: in8 = True
    if "9.0 References" in p.text: in8 = False
    if in8 and any(f"Figure {n}" in p.text for n in [38,39,40,41,42]):
        print(f"  PROBLEM: {i} {p.text[:80]}")
print("  Section 8 figure scan complete.")
