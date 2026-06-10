import sys
import re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx')

with open('scratch_matches.txt', 'w', encoding='utf-8') as out:
    out.write("=== PARAGRAPHS ===\n")
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if not txt:
            continue
        txt_lower = txt.lower()
        found = []
        if ("16" in txt_lower or "14" in txt_lower or "13" in txt_lower) and ("model" in txt_lower or "architecture" in txt_lower):
            found.append("model_count")
        if ("5" in txt_lower or "10" in txt_lower) and ("fold" in txt_lower or "cv" in txt_lower):
            found.append("fold_count")
        if "student" in txt_lower or "distill" in txt_lower:
            found.append("student/distill")
        if "logistic" in txt_lower and "regression" in txt_lower:
            found.append("logistic_regression")
            
        if found:
            out.write(f"Para {i} ({','.join(found)}):\n")
            out.write(f"{txt}\n")
            out.write("-" * 80 + "\n")
            
    out.write("\n=== TABLES ===\n")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            seen_cells = set()
            for c_idx, cell in enumerate(row.cells):
                if cell in seen_cells:
                    continue
                seen_cells.add(cell)
                txt = cell.text.strip()
                if not txt:
                    continue
                txt_lower = txt.lower()
                found = []
                if ("16" in txt_lower or "14" in txt_lower or "13" in txt_lower) and ("model" in txt_lower or "architecture" in txt_lower):
                    found.append("model_count")
                if ("5" in txt_lower or "10" in txt_lower) and ("fold" in txt_lower or "cv" in txt_lower):
                    found.append("fold_count")
                if "student" in txt_lower or "distill" in txt_lower:
                    found.append("student/distill")
                if "logistic" in txt_lower and "regression" in txt_lower:
                    found.append("logistic_regression")
                    
                if found:
                    out.write(f"Table {t_idx}, Row {r_idx}, Col {c_idx} ({','.join(found)}):\n")
                    out.write(f"{txt}\n")
                    out.write("-" * 80 + "\n")

print("Matches written to scratch_matches.txt")
