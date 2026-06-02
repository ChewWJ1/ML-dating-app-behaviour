import docx
doc = docx.Document(r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.2 (final).docx")
rels = doc.part.rels
with open("scratch/all_image_placements.txt", "w", encoding="utf-8") as f:
    for p_idx, para in enumerate(doc.paragraphs):
        blips = para._p.xpath('.//a:blip')
        if blips:
            for blip in blips:
                rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rId and rId in rels:
                    img_path = rels[rId].target_ref
                    f.write(f"P{p_idx} | {img_path}\n")
                    f.write(f"  Prev: {doc.paragraphs[p_idx-1].text.strip() if p_idx > 0 else ''}\n")
                    f.write(f"  Self: {para.text.strip()}\n")
                    f.write(f"  Next: {doc.paragraphs[p_idx+1].text.strip() if p_idx < len(doc.paragraphs)-1 else ''}\n\n")
