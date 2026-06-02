import docx

def update_explanations():
    doc_path = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx"
    doc = docx.Document(doc_path)
    
    # Specific paragraph indices found from find_outdated.py
    replacements = {
        28: [("Isotonically Calibrated Random Forest", "LightGBM (Tuned)"), ("60.48%", "57.81%")],
        240: [("Random Forest", "LightGBM (Tuned)")],
        241: [("Random Forest", "LightGBM (Tuned)"), ("9,985", "a significant portion of"), ("60.30%", "57.81%")],
        287: [("Random Forest", "LightGBM (Tuned)")],
        288: [("Random Forests", "LightGBM models"), ("Random Forest", "LightGBM (Tuned)")],
        398: [("Isotonically Calibrated Random Forest", "LightGBM (Tuned)"), ("60.48%", "57.81%"), ("Random Forest", "LightGBM (Tuned)")],
        # Update other specific texts to ensure V8 metrics
        255: [("Random Forest", "LightGBM")],
    }
    
    for idx, reps in replacements.items():
        if idx < len(doc.paragraphs):
            p = doc.paragraphs[idx]
            new_text = p.text
            for old_str, new_str in reps:
                new_text = new_text.replace(old_str, new_str)
            # Reassign text. This resets paragraph formatting (bold/italic) for this specific paragraph,
            # but it is necessary to ensure the string replacement works reliably across split runs.
            if p.text != new_text:
                p.text = new_text
                print(f"Updated paragraph {idx}")

    doc.save(doc_path)
    print("Document successfully updated with new explanations.")

if __name__ == "__main__":
    update_explanations()
