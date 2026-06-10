import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

DOCX_PATH = r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx'
doc = Document(DOCX_PATH)

print("=" * 70)
print("DISCREPANCY 1: Feature counts")
print("=" * 70)
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if any(x in t for x in ["116 input", "116 features", "67 features", "67 selected",
                              "expanding features from 25", "selected features", "features selected",
                              "feature selection union", "67 top"]):
        print(f"  Para {i}: {t}\n")

for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        seen = set()
        for ci, cell in enumerate(row.cells):
            if cell in seen: continue
            seen.add(cell)
            t = cell.text
            if any(x in t for x in ["116 input", "116 features", "67 features", "67 selected",
                                      "expanding features from 25", "features selected", "67 top"]):
                print(f"  Table {ti}, R{ri}, C{ci}: {t}\n")

print("\n" + "=" * 70)
print("DISCREPANCY 2: Engineered feature names")
print("=" * 70)
keywords = ["popularity_density", "bio_message_interaction", "selective_emoji_swiper",
            "engagement_score", "profile_completeness", "activity_intensity",
            "selectivity_ratio", "late_night_user", "engineered interaction"]
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if any(k.lower() in t.lower() for k in keywords):
        print(f"  Para {i}: {t}\n")

for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        seen = set()
        for ci, cell in enumerate(row.cells):
            if cell in seen: continue
            seen.add(cell)
            t = cell.text
            if any(k.lower() in t.lower() for k in keywords):
                print(f"  Table {ti}, R{ri}, C{ci}: {t}\n")

print("\n" + "=" * 70)
print("DISCREPANCY 3: DML p-value")
print("=" * 70)
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if any(x in t for x in ["p > 0.60", "p>0.60", "0.0322", "0.60", "indistinguishable",
                              "ATE is", "Average Treatment Effect", "causal effect", "p-value"]):
        print(f"  Para {i}: {t}\n")

for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        seen = set()
        for ci, cell in enumerate(row.cells):
            if cell in seen: continue
            seen.add(cell)
            t = cell.text
            if any(x in t for x in ["p > 0.60", "p>0.60", "0.0322", "0.60", "indistinguishable",
                                      "ATE is", "Average Treatment Effect", "causal effect"]):
                print(f"  Table {ti}, R{ri}, C{ci}: {t}\n")
