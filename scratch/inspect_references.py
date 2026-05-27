import docx
from docx import Document

doc = Document("reports/WIA1006 Machine Learning - Tying the Data Knot Assignment Report.docx")
print(f"Total paragraphs in template: {len(doc.paragraphs)}")

# Find paragraphs starting from "References" heading
found_refs = False
for idx, p in enumerate(doc.paragraphs):
    if "References" in p.text and len(p.text) < 20:
        print(f"Found References heading at index {idx}: '{p.text}'")
        found_refs = True
        # Print the next 40 paragraphs
        for next_idx in range(idx, min(idx + 50, len(doc.paragraphs))):
            print(f"{next_idx}: {doc.paragraphs[next_idx].text}")
        break

if not found_refs:
    print("Could not find 'References' heading. Printing the last 40 paragraphs instead:")
    start = max(0, len(doc.paragraphs) - 40)
    for next_idx in range(start, len(doc.paragraphs)):
         print(f"{next_idx}: {doc.paragraphs[next_idx].text}")
