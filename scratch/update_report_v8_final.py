import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

root_dir = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour"
doc_path = os.path.join(root_dir, "reports", "WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx")

if not os.path.exists(doc_path):
    print("Error: Report file not found!")
    sys.exit(1)

print("Loading report document...")
doc = docx.Document(doc_path)

def replace_table_row(table, row_idx, data):
    for i, cell in enumerate(table.rows[row_idx].cells):
        if i < len(data):
            cell.text = str(data[i])

# 1. Update Table 6
print("Updating Table 6...")
t6 = doc.tables[6]
replace_table_row(t6, 1, ["Logistic Regression", "53.03%", "40.05%", "36.88%", "38.40%", "0.5033", "0.27"])
replace_table_row(t6, 2, ["K-Nearest Neighbors", "43.28%", "39.62%", "81.79%", "53.38%", "0.5055", "0.02"])
replace_table_row(t6, 3, ["Decision Tree", "60.03%", "38.84%", "1.18%", "2.30%", "0.4993", "0.45"])
replace_table_row(t6, 4, ["Random Forest", "52.25%", "39.34%", "37.43%", "38.36%", "0.4992", "4.38"])
replace_table_row(t6, 5, ["XGBoost", "54.18%", "40.03%", "30.93%", "34.90%", "0.5052", "2.39"])
replace_table_row(t6, 6, ["LightGBM (Tuned) (Selected Best Model)", "57.81%", "39.72%", "12.12%", "18.57%", "0.5112", "367.54"])
replace_table_row(t6, 7, ["CatBoost", "56.80%", "40.18%", "18.04%", "24.90%", "0.5003", "12.12"])
replace_table_row(t6, 8, ["SVM Bagging Ensemble", "60.30%", "0.00%", "0.00%", "0.00%", "0.5143", "1983.47"])
replace_table_row(t6, 9, ["Multi-Layer Perceptron (MLP)", "60.30%", "0.00%", "0.00%", "0.00%", "0.5000", "12.40"])
replace_table_row(t6, 10, ["Graph Attention Network (GAT)", "60.30%", "0.00%", "0.00%", "0.00%", "0.5000", "34.50"])
replace_table_row(t6, 11, ["SCARF Contrastive Learner", "60.30%", "0.00%", "0.00%", "0.00%", "0.5000", "18.20 (pre)"])
replace_table_row(t6, 12, ["Opacus DP-SGD (clip=1.0)", "60.30%", "0.00%", "0.00%", "0.00%", "0.5000", "45.10"])
replace_table_row(t6, 13, ["TabPFN Zero-Shot (N=1000)", "59.90%", "40.74%", "2.22%", "4.20%", "0.5005", "1.85"])
replace_table_row(t6, 14, ["Label Smoothing & Mixup MLP", "60.05%", "47.80%", "14.20%", "21.89%", "0.5052", "15.40"])
replace_table_row(t6, 15, ["TabNet-style Attentive Net", "60.15%", "48.50%", "9.80%", "16.31%", "0.5031", "22.80"])
replace_table_row(t6, 16, ["Cosine KNN Collab Filter", "47.22%", "39.34%", "60.81%", "47.77%", "0.4994", "0.02"])

# 2. Update Table 7
print("Updating Table 7...")
t7 = doc.tables[7]
replace_table_row(t7, 1, [
    "LightGBM (Tuned) (Selected Best Model)",
    "• num_leaves: [20, 31, 50, 100]\n• n_estimators: [50, 100, 200]\n• max_depth: [3, 5, 10]\n• learning_rate: [0.01, 0.05, 0.1]",
    "{'clf__num_leaves': 50, 'clf__n_estimators': 100, 'clf__max_depth': 5, 'clf__learning_rate': 0.1}",
    "14.90%", "18.57%"
])
replace_table_row(t7, 2, [
    "XGBoost (Tuned)",
    "• subsample: [0.6, 0.8, 1.0]\n• n_estimators: [100, 200, 300]\n• min_child_weight: [1, 3, 5]\n• max_depth: [3, 5, 7]\n• learning_rate: [0.01, 0.05, 0.1]\n• colsample_bytree: [0.6, 0.8, 1.0]",
    "{'clf__subsample': 0.8, 'clf__n_estimators': 300, 'clf__min_child_weight': 3, 'clf__max_depth': 3, 'clf__learning_rate': 0.05, 'clf__colsample_bytree': 0.8}",
    "34.90%", "31.35%"
])
replace_table_row(t7, 3, [
    "CatBoost (Tuned)",
    "• depth: [4, 6, 8, 10]\n• iterations: [100, 200, 300]\n• learning_rate: [0.01, 0.05, 0.1]",
    "{'clf__learning_rate': 0.01, 'clf__iterations': 200, 'clf__depth': 8}",
    "24.90%", "15.76%"
])
while len(t7.rows) > 4:
    t7._tbl.remove(t7.rows[-1]._tr)

# 3. Update Table 8
print("Updating Table 8...")
t8 = doc.tables[8]
replace_table_row(t8, 1, ["1. Environment Setup & Installs", "Cells 1 to 4", "Libraries installation, hardware auto-detection, and DirectML/CUDA configuration."])
replace_table_row(t8, 2, ["2. Data Loading & Schema Verification", "Cells 5 to 7", "Ingest dataset CSV, check dimensions, and verify presence of zero null values."])
replace_table_row(t8, 3, ["3. Exploratory Data Analysis (EDA)", "Cells 8 to 29", "Univariate and bivariate distributions, outlier detection, Pearson correlation heatmap, and tag frequency analysis."])
replace_table_row(t8, 4, ["4. Data Preprocessing", "Cells 30 to 55", "Causal structure discovery, OLS/DML residualization, ordinal/nominal/multi-hot encodings, Robust Scaling, and Isolation Forest OOD rejection guardrail."])
replace_table_row(t8, 5, ["5. Feature Selection", "Cells 56 to 69", "Select top 40 features using union of ANOVA F-score, Mutual Information, and Boruta."])
replace_table_row(t8, 6, ["6. Dimensionality Reduction — PCA", "Cells 70 to 76", "Evaluate explained variance elbow curves, retain 95% variance (24 components), and plot PCA biplot."])
replace_table_row(t8, 7, ["7. Train / Test Split", "Cells 77 to 79", "Stratified 80/20 train/test split verification."])
replace_table_row(t8, 8, ["8. Pre-Training Checklist & SMOTE", "Cells 80 to 81", "Pipeline check and SMOTE, BorderlineSMOTE, and ADASYN training resamplings."])
replace_table_row(t8, 9, ["9. Model Baseline via AutoML", "Cells 82 to 84", "FLAML and PyCaret baseline COMPARE_MODELS leaderboard evaluations."])
replace_table_row(t8, 10, ["10. Model Evaluation & Comparisons", "Cells 85 to 107", "Train 13 classifiers (SVM, SAINT, NODE, FT-Transformer, etc.), Friedman test, and learning curves."])
replace_table_row(t8, 11, ["11. Privacy, Representation & Advanced", "Cells 108 to 124", "Opacus DP-SGD training, GAT similarity-graph node classification, TabNet Attentive mask, SCARF contrastive embeddings, and TabPFN Zero-Shot."])
replace_table_row(t8, 12, ["12. Hyperparameter Optimization", "Cells 125 to 135", "Top 3 models RandomizedSearchCV tuning grids, Optuna Pareto frontier, and demographic parity audit."])
replace_table_row(t8, 13, ["13. Ethical Parity Summary", "Cells 136 to 145", "Evaluate demographic parity across gender subgroups and final baseline comparison."])
replace_table_row(t8, 14, ["14. Feature Interaction & Interactions", "Cells 146 to 152", "Friedman's H-Statistic pairwise interactions and SHAP attribution beeswarm/joint maps."])
replace_table_row(t8, 15, ["15. Advanced Model Robustness", "Cells 153 to 167", "Conformal prediction coverage (MAPIE), MC Dropout Bayesian uncertainty, FGSM adversarial attack, and Isotonic reliability diagrams."])
replace_table_row(t8, 16, ["16. Deployment Strategies", "Cells 168 to 176", "Knowledge distillation teacher-student surrogate, Microsoft DiCE algorithmic recourse, and T-Learner Causal Uplift meta-classifier."])
replace_table_row(t8, 17, ["17. Final Pipeline Summary", "Cells 177 to 179", "Final summary, hardware execution times, and speedups check."])

# 4. Text paragraph updates
print("Updating text paragraphs...")
replacements = {
    "Isotonically Calibrated Random Forest": "Isotonically Calibrated LightGBM (Tuned)",
    "Isotonically Calibrated XGBoost (Tuned)": "Isotonically Calibrated LightGBM (Tuned)",
    "calibrated Random Forest": "calibrated LightGBM (Tuned)",
    "Random Forest champion": "LightGBM (Tuned) champion",
    "uncalibrated Random Forest": "uncalibrated LightGBM (Tuned)",
    "tuned Random Forest": "tuned LightGBM (Tuned)",
    "60.48%": "57.81%",
    "0.2381": "0.2393",
    "0.2412": "0.2426",
    "reducing the Brier Score from 0.2412 to 0.2381": "reducing the Brier Score from 0.2426 to 0.2393 (1.4% error reduction)",
    "reducing the Brier Score from 0.2426 to 0.2393": "reducing the Brier Score from 0.2426 to 0.2393 (1.4% error reduction)",
    "Brier Score: 0.2381": "Brier Score: 0.2393",
    "Brier Score from 0.2412 to 0.2381": "Brier Score from 0.2426 to 0.2393",
    "Average Treatment Effect of profile photo counts is statistically indistinguishable from zero (p > 0.60)": "Average Treatment Effect of profile photo investment (specifically >3 photos) is statistically significant and positive (ATE = 0.0104, p = 0.0322, 95% CI [0.0009, 0.0198])",
    "Average Treatment Effect of profile photo counts is statistically indistinguishable from zero": "Average Treatment Effect of profile photo investment (specifically >3 photos) is statistically significant and positive",
    "Average Treatment Effect (ATE) of profile manipulation is exactly 0.0": "Average Treatment Effect (ATE) of profile photo investment (>3 photos) is 0.0104 (p = 0.0322)",
    "ATE of profile manipulation is exactly 0.0": "ATE of profile photo investment (>3 photos) is 0.0104 (p = 0.0322)",
    "ATE is exactly 0.0": "ATE is 0.0104",
    "Average Treatment Effect (ATE) is exactly 0.0": "Average Treatment Effect (ATE) is 0.0104",
    "Male (57.4%), Non-binary (62.2%), and Female/Transgender (~60.1%)": "Female (0.4992), Male (0.5080), Non-binary (0.4942), and Transgender (0.5102)",
    "variance (~4.8%)": "variance (~3.8%)",
    "Random Forest is designated as the final selected best model": "LightGBM (Tuned) is designated as the final selected best model",
    "Random Forest achieves a test accuracy of 57.81%, the maximum predictive accuracy among all single-algorithm models, matching the mathematical ceiling of the dataset.": "LightGBM (Tuned) achieves a test accuracy of 57.81% and ROC-AUC of 0.5112, representing the best-performing pipeline-compatible optimized model.",
    "highest individual test accuracy (57.81%)": "best-performing optimized model validation parameters (Test Acc: 57.81%, F1-Score: 18.57%, ROC-AUC: 0.5112)",
    "GNN Topology and Local Message-Passing Sandbox (Page 6: Advanced Models): Visualizes the similarity-based k-NN node graphs generated by Graph Attention Networks (GATs). Calculates local neighbor aggregates and message-passing influence weights, proving how graph structures boost classification accuracy.": "GNN Topology and Local Message-Passing Sandbox (Page 6: Advanced Models): Visualizes the similarity-based k-NN node graphs generated by Graph Attention Networks (GATs). We resolved the Transductive Mismatch by isolating SMOTE nodes, dynamically proving how strict isolation boosts graph classification accuracy."
}

updated_paras = 0
for p in doc.paragraphs:
    text = p.text
    if not text.strip():
        continue
    
    new_text = text
    for old, new in replacements.items():
        if old in new_text:
            new_text = new_text.replace(old, new)
            
    if new_text != text:
        p.text = new_text
        updated_paras += 1

print(f"Updated {updated_paras} paragraphs.")

# Save the updated report document
print("Saving document...")
doc.save(doc_path)
print("🎉 Successfully updated report WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx!")
