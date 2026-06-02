import docx
import sys
import os

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.2 (final).docx"
doc = docx.Document(doc_path)

print(f"Total tables: {len(doc.tables)}")
for idx, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns) if rows > 0 else 0
    print(f"\nTable {idx}: Rows={rows}, Cols={cols}")
    if rows > 0:
        first_row_text = []
        for cell in table.rows[0].cells:
            # avoid duplicates from merged cells
            txt = cell.text.strip().replace('\n', ' ')
            if not first_row_text or first_row_text[-1] != txt:
                first_row_text.append(txt)
        print(f"  Header cells: {first_row_text}")
        if rows > 1:
            second_row_text = []
            for cell in table.rows[1].cells:
                txt = cell.text.strip().replace('\n', ' ')
                if not second_row_text or second_row_text[-1] != txt:
                    second_row_text.append(txt)
            print(f"  Second row: {second_row_text}")
