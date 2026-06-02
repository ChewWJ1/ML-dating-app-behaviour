import docx
import sys
import re

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.2 (final).docx"
doc = docx.Document(doc_path)

rels = doc.part.rels

print("--- Paragraph Images ---")
for p_idx, para in enumerate(doc.paragraphs):
    blips = para._p.xpath('.//a:blip')
    if blips:
        image_paths = []
        for blip in blips:
            rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rId and rId in rels:
                image_paths.append(rels[rId].target_ref)
        if image_paths:
            # Let's find surrounding figure caption
            caption = ""
            for i in range(max(0, p_idx - 3), min(len(doc.paragraphs), p_idx + 4)):
                txt = doc.paragraphs[i].text.strip()
                if txt.startswith("Figure ") or txt.startswith("*Figure "):
                    caption = txt
                    break
            print(f"P{p_idx}: Caption='{caption}' -> Images: {image_paths}")
