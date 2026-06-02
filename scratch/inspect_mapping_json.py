import docx
import json
import re

doc_path = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.2 (final).docx"
doc = docx.Document(doc_path)

rels = doc.part.rels
rId_to_image = {}
for rId, rel in rels.items():
    if "image" in rel.target_ref:
        rId_to_image[rId] = rel.target_ref

mapping = []
for p_idx, para in enumerate(doc.paragraphs):
    xml_str = para._p.xml
    rIds = re.findall(r'r:embed="([^"]+)"', xml_str)
    rIds.extend(re.findall(r'r:link="([^"]+)"', xml_str))
    if rIds:
        unique_rIds = list(set(rIds))
        resolved = [rId_to_image.get(rId) for rId in unique_rIds if rId in rId_to_image]
        if resolved:
            # Let's find surrounding figure caption
            caption = ""
            for i in range(max(0, p_idx - 3), min(len(doc.paragraphs), p_idx + 4)):
                txt = doc.paragraphs[i].text.strip()
                if txt.startswith("Figure ") or txt.startswith("*Figure "):
                    caption = txt
                    break
            mapping.append({
                "p_idx": p_idx,
                "text": para.text.strip(),
                "images": resolved,
                "caption": caption
            })

with open("scratch/image_mapping.json", "w", encoding="utf-8") as f:
    json.dump(mapping, f, indent=2, ensure_ascii=False)
print("Saved mapping of", len(mapping), "images.")
