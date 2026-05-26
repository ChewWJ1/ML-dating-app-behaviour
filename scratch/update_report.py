import docx

doc_path = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report.docx"
doc = docx.Document(doc_path)

# 1. Insert descriptions in Section 3.3
# Let's search for the paragraph containing "Champion Stacking Ensemble:"
target_para_idx = -1
for idx, para in enumerate(doc.paragraphs):
    if "Champion Stacking Ensemble:" in para.text:
        target_para_idx = idx
        break

if target_para_idx != -1:
    print(f"Found target paragraph at index {target_para_idx}")
    
    # We will insert paragraphs after this target paragraph
    # To insert a paragraph at a specific position using python-docx, we use paragraph._element.addnext()
    # Or we can just insert them by accessing the parent element.
    # Let's define the new texts:
    new_descriptions = [
        "FT-Transformer (Feature Tokenizer Transformer): Projects both numerical and categorical features into dense token embeddings using projection linear layers and embedding dictionaries, then processes them through multi-head self-attention blocks to capture complex cross-feature interactions.",
        "SAINT (Self-Attention and Invariant Representation): Applies self-attention over the feature dimensions (across features) rather than just spatial dimensions, capturing non-linear relationships and cross-feature correlations on tabular inputs.",
        "NODE (Neural Oblivious Decision Ensembles): Integrates differentiable oblivious decision trees with entmax/softmax activations, allowing forest-based tabular networks to be trained natively via backpropagation on GPUs."
    ]
    
    current_p = doc.paragraphs[target_para_idx]
    for text in new_descriptions:
        # insert_paragraph_after helper
        new_p = doc.add_paragraph()
        new_p.text = text
        # Move the new paragraph to be right after current_p in the XML tree
        current_p._element.addnext(new_p._element)
        current_p = new_p
    print("Inserted new model descriptions.")
else:
    print("Champion Stacking Ensemble paragraph not found.")

# 2. Add rows to Table 9 (Optimizations Table)
table = doc.tables[9]
print(f"Updating Table 9 (current rows: {len(table.rows)})")

new_rows_data = [
    (
        "Dynamic Hardware Auto-Detection Engine",
        "Compute & Cross-Device Execution",
        "Pipeline execution crashes or slows down when running on different GPUs or fallback devices across teammate environments.",
        "Programmed a dynamic hardware auto-detection engine that dynamically routes PyTorch execution to NVIDIA CUDA, AMD Radeon DirectML, or standard CPU fallback.",
        "Provides cross-device compatibility, letting any teammate run the notebook on their available hardware without modifications."
    ),
    (
        "Custom PyTorch Sklearn Wrapper",
        "Workflow Compatibility",
        "Custom PyTorch architectures cannot natively run inside scikit-learn cross-validation, metric generators, or parameter tuning loops.",
        "Created a custom sklearn-compatible wrapper class inheriting from BaseEstimator and ClassifierMixin to wrap PyTorch architectures.",
        "Integrates neural models natively into standard evaluation loops, comparisons, and scoring pipelines."
    ),
    (
        "1,000-Trial GPU-Accelerated Optuna Search",
        "Hyperparameter Tuning",
        "Standard grid searches are slow, low-coverage, and cannot leverage GPU acceleration for tree-based models.",
        "Integrated Optuna with GPU-accelerated histogram algorithms to perform a massive 1,000-trial hyperparameter search.",
        "Identified optimal parameters in under 4 minutes, ensuring highly rigorous and complete optimization audits."
    ),
    (
        "SVM-only Bypass and models_advanced Routing",
        "Workflow Efficiency",
        "Retraining the RBF SVM takes 30+ minutes, causing massive delays when retraining other models from scratch.",
        "Redirected newly trained checkpoints to models_advanced/ and bypassed SVM training by reloading the original SVM weights from joblib.",
        "Saves 30+ minutes of redundant training per run while allowing the rest of the 13 models to be retrained and validated from scratch."
    )
]

for data in new_rows_data:
    row = table.add_row()
    for col_idx, text in enumerate(data):
        row.cells[col_idx].text = text
print(f"Table 9 updated (new rows: {len(table.rows)})")

# Save document
doc.save(doc_path)
print("Report updated and saved successfully.")
