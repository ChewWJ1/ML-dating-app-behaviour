import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

root_dir = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour"
doc_path = os.path.join(root_dir, "reports", "WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx")

doc = docx.Document(doc_path)
print(f"Total tables: {len(doc.tables)}")

for idx, table in enumerate(doc.tables):
    first_row = [c.text.strip().replace('\n', ' ') for c in table.rows[0].cells]
    first_row_str = " | ".join(first_row)[:120]
    print(f"Table {idx:2d}: {first_row_str}")
