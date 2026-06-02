import docx
import json
import os
import re

doc_path = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.2 (final).docx"
doc = docx.Document(doc_path)
rels = doc.part.rels

mapping = {}
for p_idx, para in enumerate(doc.paragraphs):
    blips = para._p.xpath('.//a:blip')
    if blips:
        for blip in blips:
            rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rId and rId in rels:
                img_path = rels[rId].target_ref
                # Find caption
                caption = ""
                # Search up to 3 paragraphs before/after
                for i in range(max(0, p_idx - 3), min(len(doc.paragraphs), p_idx + 4)):
                    txt = doc.paragraphs[i].text.strip()
                    if txt.startswith("Figure ") or txt.startswith("*Figure ") or txt.startswith("Figure:"):
                        caption = txt
                        break
                
                # Clean path
                img_name = os.path.basename(img_path)
                mapping[img_name] = {
                    "p_idx": p_idx,
                    "caption": caption,
                    "target_ref": img_path
                }

with open("scratch/full_image_mapping.json", "w", encoding="utf-8") as f:
    json.dump(mapping, f, indent=2)

print(f"Mapped {len(mapping)} images:")
for name, val in sorted(mapping.items(), key=lambda x: int(re.search(r'\d+', x[0]).group()) if re.search(r'\d+', x[0]) else 0):
    print(f"  {name} -> {val['caption']}")
