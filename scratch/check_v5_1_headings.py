import docx
import os

doc_path = r"reports/WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.1(long).docx"
if not os.path.exists(doc_path):
    print("Error: File not found!")
    exit(1)

doc = docx.Document(doc_path)
print(f"Total paragraphs: {len(doc.paragraphs)}")

for idx in range(290, len(doc.paragraphs)):
    text = doc.paragraphs[idx].text.strip()
    if text.startswith(("1.0", "2.0", "3.0", "4.0", "5.0", "6.0", "7.0", "8.0", "9.0")) or "Figure " in text:
        print(f"P {idx} ({doc.paragraphs[idx].style.name}): {text}")
