import sys
import re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx')

print("=== SEARCH FOR 'finding' OR 'fold' IN SECTION 17 / PARAS 385-425 ===")
for i in range(380, min(426, len(doc.paragraphs))):
    txt = doc.paragraphs[i].text
    txt_lower = txt.lower()
    if "finding" in txt_lower or "fold" in txt_lower or "cv" in txt_lower or "cross-validation" in txt_lower:
        print(f"Para {i}: {txt}")
        print("-" * 60)
        
print("=== SEARCH FOR ALL OCCURRENCES OF '5-fold' or '5 fold' or '10-fold' or '10 fold' ===")
for i, p in enumerate(doc.paragraphs):
    txt = p.text
    if re.search(r'\b(5|10)-?folds?\b', txt, re.IGNORECASE):
        print(f"Para {i}: {txt}")
        print("-" * 60)
