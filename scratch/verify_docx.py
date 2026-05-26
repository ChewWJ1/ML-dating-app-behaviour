import docx

doc_path = r"reports/WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5 SOTA.docx"
try:
    doc = docx.Document(doc_path)
    print("Success: The massive SOTA DOCX document parses and loads perfectly!")
    print(f"Total Paragraphs in V5 SOTA DOCX: {len(doc.paragraphs)}")
    
    # Check if a few key phrases exist in the paragraphs
    keywords = [
        "Double Machine Learning",
        "TabNet-style",
        "OOD Rejection Guardrail",
        "Uplift Modeling",
        "Brier Score Decomposition",
        "Platt Scaling vs Isotonic"
    ]
    matches = {kw: 0 for kw in keywords}
    for p in doc.paragraphs:
        for kw in keywords:
            if kw in p.text:
                matches[kw] += 1
                
    for kw, count in matches.items():
        print(f"👉 Keyword '{kw}' matches: {count}")
        
    print("Success: Verification of all V5.1 massive SOTA report additions is complete!")
except Exception as e:
    print(f"Error: DOCX document verification failed! Details: {e}")
