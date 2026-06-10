import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx')

paras_to_print = [27, 28, 57, 67, 163, 183, 197, 211, 212, 240, 281, 337, 341, 367, 379, 395, 410, 411]

for p_idx in paras_to_print:
    if p_idx < len(doc.paragraphs):
        print(f"Para {p_idx}:")
        print(doc.paragraphs[p_idx].text)
        print("=" * 60)
