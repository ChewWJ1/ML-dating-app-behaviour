import docx

doc_long = docx.Document(r"reports/WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.1(long).docx")
doc_sota = docx.Document(r"reports/WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.2 SOTA.docx")

def get_fig_counts(doc):
    fig_counts = {}
    captions = []
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if "Figure " in text:
            captions.append((idx, text))
            parts = text.split(":")
            if parts:
                label = parts[0].strip()
                if label.startswith("Figure"):
                    fig_counts[label] = fig_counts.get(label, 0) + 1
    return fig_counts, captions

print("\n--- Long Report Figure Counts ---")
long_counts, long_caps = get_fig_counts(doc_long)
dup_long = {k: v for k, v in long_counts.items() if v > 1}
print(f"Duplicates in Long: {dup_long}")
for label, count in dup_long.items():
    print(f"Details for {label}:")
    for idx, text in long_caps:
        if text.startswith(label):
            print(f"  P {idx}: {text}")

print("\n--- SOTA Report Figure Counts ---")
sota_counts, sota_caps = get_fig_counts(doc_sota)
dup_sota = {k: v for k, v in sota_counts.items() if v > 1}
print(f"Duplicates in SOTA: {dup_sota}")
for label, count in dup_sota.items():
    print(f"Details for {label}:")
    for idx, text in sota_caps:
        if text.startswith(label):
            print(f"  P {idx}: {text}")
