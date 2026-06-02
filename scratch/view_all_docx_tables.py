import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx"
doc = docx.Document(doc_path)

print(f"Total tables: {len(doc.tables)}")

for idx, table in enumerate(doc.tables):
    print(f"\n==========================================")
    print(f"TABLE {idx}")
    print(f"==========================================")
    for r_idx, row in enumerate(table.rows):
        row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        print(f"Row {r_idx}: {' | '.join(row_cells)}")
