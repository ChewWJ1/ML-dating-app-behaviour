import docx

doc_path = "reports/WIA1006 Machine Learning - Tying the Data Knot Assignment Report.docx"
doc = docx.Document(doc_path)

print(f"Total Paragraphs: {len(doc.paragraphs)}")
print(f"Total Tables: {len(doc.tables)}")

# Let's inspect some paragraphs and print them
print("\n--- First 30 Paragraphs of Original ---")
for i in range(min(50, len(doc.paragraphs))):
    p = doc.paragraphs[i]
    if p.text.strip():
        print(f"P {i} ({p.style.name}): {p.text[:120]}")

print("\n--- Outline of Headings & Sections ---")
for idx, p in enumerate(doc.paragraphs):
    if p.style.name.startswith("Heading") or (p.text.strip() and p.text.strip()[0].isdigit() and p.text.split()[0].replace('.', '').isdigit()):
        print(f"P {idx} ({p.style.name}): {p.text[:100]}")
