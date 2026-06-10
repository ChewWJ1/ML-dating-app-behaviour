import sys
import re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx')

def check_text(text, source_info):
    text_lower = text.lower()
    # Check for number 16 or 14 and 'model' or 'architecture'
    # Check for '5' or '10' and 'fold' or 'cv'
    # Check for 'student' or 'distill'
    # Check for 'logistic' and 'regression' or 'student'
    found = []
    if ("16" in text_lower or "14" in text_lower or "13" in text_lower) and ("model" in text_lower or "architecture" in text_lower):
        found.append("model_count")
    if ("5" in text_lower or "10" in text_lower) and ("fold" in text_lower or "cv" in text_lower):
        found.append("fold_count")
    if "student" in text_lower or "distill" in text_lower:
        found.append("student/distill")
    if "logistic" in text_lower and "regression" in text_lower:
        found.append("logistic_regression")
        
    if found:
        print(f"[{','.join(found)}] in {source_info}:")
        print(f"  {text}")
        print("=" * 60)

print("=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        check_text(p.text, f"Para {i}")

print("=== TABLES ===")
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        seen_cells = set()
        for c_idx, cell in enumerate(row.cells):
            if cell in seen_cells:
                continue
            seen_cells.add(cell)
            if cell.text.strip():
                check_text(cell.text, f"Table {t_idx}, Row {r_idx}, Col {c_idx}")
