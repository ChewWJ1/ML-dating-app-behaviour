import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

DOCX_PATH = r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx'
doc = Document(DOCX_PATH)

p = doc.paragraphs[169]
full = ''.join(r.text for r in p.runs)
new = full.replace('67-dimensional', '66-dimensional', 1)
if p.runs:
    p.runs[0].text = new
    for r in p.runs[1:]: r.text = ''

doc.save(DOCX_PATH)
print('Fixed and saved.')
print('Para 169 now reads:', doc.paragraphs[169].text[:120])
