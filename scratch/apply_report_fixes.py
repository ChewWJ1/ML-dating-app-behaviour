"""
Script to apply targeted report fixes to WIA1006 ML Report V8 (final).docx:
1. Model count corrections (16/14 models -> 13 models)
2. CV fold count corrections (5-fold -> 10-fold)
3. Student description correction (logistic regression -> StudentNet)
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

DOCX_PATH = r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx'
doc = Document(DOCX_PATH)

changes_made = []

def replace_in_para(para, old_text, new_text, source_info):
    """Replace text in a paragraph across runs, preserving style/formatting structure."""
    full_text = ''.join(r.text for r in para.runs)
    if old_text not in full_text:
        return False
    
    new_full_text = full_text.replace(old_text, new_text, 1)
    
    # Clear runs and place updated text in first run
    if para.runs:
        para.runs[0].text = new_full_text
        for run in para.runs[1:]:
            run.text = ''
    else:
        para.text = new_full_text
        
    changes_made.append(f"{source_info}: '{old_text[:50]}...' -> '{new_text[:50]}...'")
    return True

# ============================================================
# FIX 1: Model Count Corrections
# ============================================================
print("Applying model count corrections...")

# Para 27
p27 = doc.paragraphs[27]
old_27 = "14 custom architectures and 2 AutoML baselines (total 16 models)"
new_27 = "11 custom architectures and 2 AutoML baselines (total 13 models)"
replace_in_para(p27, old_27, new_27, "Para 27")

# Para 163
p163 = doc.paragraphs[163]
old_163 = "We evaluated 16 diverse models to establish performance baselines."
new_163 = "We evaluated 13 diverse models to establish performance baselines."
replace_in_para(p163, old_163, new_163, "Para 163")

# Para 183
p183 = doc.paragraphs[183]
old_183 = "all 14 baseline and advanced models are tabulated below:"
new_183 = "all 13 baseline and advanced models are tabulated below:"
replace_in_para(p183, old_183, new_183, "Para 183")

# Para 395
p395 = doc.paragraphs[395]
old_395 = "We evaluated 16 baseline and advanced classifiers (Logistic Regression, KNN, DT, RF, XGBoost, custom multi-threaded Bagging SVM, LightGBM, CatBoost, MLP, Label Smoothing MLP, Cosine KNN CF, GAT Graph Neural Network, SCARF Contrastive Learner, Opacus DP-SGD, TabPFN Zero-Shot, and TabNet-style Attentive Net)"
new_395 = "We evaluated 13 baseline and advanced classifiers (Logistic Regression, KNN, Decision Tree, Random Forest, XGBoost, custom multi-threaded Bagging SVM, LightGBM, CatBoost, Balanced Random Forest, KNN (Cosine Metric), FT-Transformer, SAINT, and NODE)"
replace_in_para(p395, old_395, new_395, "Para 395")


# ============================================================
# FIX 2: Cross-Validation Fold Count Corrections
# ============================================================
print("Applying cross-validation fold count corrections...")

# Para 211
p211 = doc.paragraphs[211]
old_211 = "Figure 24: 5-Fold Cross-Validation"
new_211 = "Figure 24: 10-Fold Cross-Validation"
replace_in_para(p211, old_211, new_211, "Para 211")

# Para 212
p212 = doc.paragraphs[212]
old_212 = "presents the 5-fold cross-validation accuracies"
new_212 = "presents the 10-fold cross-validation accuracies"
replace_in_para(p212, old_212, new_212, "Para 212")

# Para 240
p240 = doc.paragraphs[240]
old_240 = "using a 5-fold CV RandomizedSearchCV."
new_240 = "using a 10-fold CV RandomizedSearchCV."
replace_in_para(p240, old_240, new_240, "Para 240")

# Para 367
p367 = doc.paragraphs[367]
old_367_1 = "across 5 independent validation folds."
new_367_1 = "across 10 independent validation folds."
old_367_2 = "statistically weak 5-fold paired t-test"
new_367_2 = "statistically weak 10-fold paired t-test"
replace_in_para(p367, old_367_1, new_367_1, "Para 367 (folds)")
replace_in_para(p367, old_367_2, new_367_2, "Para 367 (t-test)")

# Para 379
p379 = doc.paragraphs[379]
old_379 = "distribute the 5 validation folds cleanly"
new_379 = "distribute the 10 validation folds cleanly"
replace_in_para(p379, old_379, new_379, "Para 379")

# Table 4, Row 6, Col 2
t4 = doc.tables[4]
cell_4_6_2 = t4.rows[6].cells[2]
for p in cell_4_6_2.paragraphs:
    replace_in_para(p, "5-fold CV", "10-fold CV", "Table 4, R6, C2")

# Table 9, Row 2, Col 2
t9 = doc.tables[9]
cell_9_2_2 = t9.rows[2].cells[2]
for p in cell_9_2_2.paragraphs:
    replace_in_para(p, "5-fold cross-validation", "10-fold cross-validation", "Table 9, R2, C2")

# Table 10, Row 2, Col 3
t10 = doc.tables[10]
cell_10_2_3 = t10.rows[2].cells[3]
for p in cell_10_2_3.paragraphs:
    replace_in_para(p, "5-fold cross-validation", "10-fold cross-validation", "Table 10, R2, C3")


# ============================================================
# FIX 3: Student Description Correction
# ============================================================
print("Applying student description correction...")

# Para 411
p411 = doc.paragraphs[411]
old_411 = "compressed student logistic regression model."
new_411 = "compressed student lightweight neural network (StudentNet) model."
replace_in_para(p411, old_411, new_411, "Para 411")


# ============================================================
# SAVE DOCUMENT
# ============================================================
doc.save(DOCX_PATH)

print()
print(f"=== FIXES APPLIED SUCCESSFULLY: {len(changes_made)} ===")
for c in changes_made:
    print(f"  ✓ {c}")
print()
print("Report document saved successfully.")
