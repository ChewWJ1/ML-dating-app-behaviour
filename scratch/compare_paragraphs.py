import docx
import os

doc_long = docx.Document(r"reports/WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.1(long).docx")
doc_sota = docx.Document(r"reports/WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.1 SOTA.docx")

print(f"Long report paragraphs: {len(doc_long.paragraphs)}")
print(f"SOTA report paragraphs: {len(doc_sota.paragraphs)}")

# Let's check paragraph texts at a few standard positions (e.g. 50, 100, 150)
for idx in [50, 100, 150, 220]:
    txt_l = doc_long.paragraphs[idx].text if idx < len(doc_long.paragraphs) else "N/A"
    txt_s = doc_sota.paragraphs[idx].text if idx < len(doc_sota.paragraphs) else "N/A"
    print(f"\n--- P {idx} ---")
    print(f"Long: {repr(txt_l[:100])}")
    print(f"SOTA: {repr(txt_s[:100])}")
    if txt_l == txt_s:
        print("MATCH!")
    else:
        print("DIFFERENT!")
