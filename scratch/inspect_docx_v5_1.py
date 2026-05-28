import docx
import os

doc_path = r"reports/WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.1(long).docx"
if not os.path.exists(doc_path):
    print("Error: File not found!")
    exit(1)

doc = docx.Document(doc_path)
print(f"Total paragraphs: {len(doc.paragraphs)}")

found_5_4 = False
count = 0
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "5.4" in text or "5.5" in text:
        found_5_4 = True
    if found_5_4:
        print(f"P {idx} ({p.style.name}): {text}")
        count += 1
        if count > 50:
            break
