import docx

doc = docx.Document(r"reports/WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.1 SOTA.docx")
print(f"Total paragraphs: {len(doc.paragraphs)}")

count = 0
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "Figure " in text:
        count += 1
        print(f"{count}. P {idx}: {text}")
        if count >= 30:
            break
