import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.2 (final).docx"
doc = docx.Document(doc_path)
rels = doc.part.rels

# Find all images in paragraphs and their exact text
print("--- Detailed Image Placement in Paragraphs ---")
for p_idx, para in enumerate(doc.paragraphs):
    blips = para._p.xpath('.//a:blip')
    if blips:
        for blip in blips:
            rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rId and rId in rels:
                img_path = rels[rId].target_ref
                print(f"\nP{p_idx} has image: {img_path}")
                print(f"  Para text: '{para.text.strip()}'")
                if p_idx > 0:
                    print(f"  Prev P{p_idx-1}: '{doc.paragraphs[p_idx-1].text.strip()}'")
                if p_idx < len(doc.paragraphs) - 1:
                    print(f"  Next P{p_idx+1}: '{doc.paragraphs[p_idx+1].text.strip()}'")
