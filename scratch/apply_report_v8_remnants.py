import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

root_dir = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour"
doc_path = os.path.join(root_dir, "reports", "WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx")

if not os.path.exists(doc_path):
    print("Error: Report file not found!")
    sys.exit(1)

doc = docx.Document(doc_path)

# Update Paragraph 201
p201 = doc.paragraphs[201]
print("Old Paragraph 201:", p201.text)
p201.text = "Based on the comprehensive metrics in Table 6, LightGBM (Tuned) is designated as the final selected best model for the following reasons:"
print("New Paragraph 201:", p201.text)

# Update Paragraph 206
p206 = doc.paragraphs[206]
print("Old Paragraph 206:", p206.text)
p206.text = "Regularized Gradient Boosting Optimization: By sequentially minimizing residual errors through regularized gradient boosting with a leaf-wise growth strategy, the Champion Model stabilizes predictions and controls overfitting, making it the most mathematically robust single classifier against the high noise levels present in this synthetic dataset."
print("New Paragraph 206:", p206.text)

# Update Table 10, Row 3, Cell 3
t10 = doc.tables[10]
cell = t10.rows[3].cells[3]
print("Old Table 10 Cell:", cell.text)
cell.text = "Deployed SHAP (Shapley Additive exPlanations) TreeExplainer on the selected best model (LightGBM (Tuned)) to analyze attributions."
print("New Table 10 Cell:", cell.text)

# Save the document
print("Saving report docx...")
doc.save(doc_path)
print("🎉 Successfully applied remaining V8 champion model updates to the DOCX report!")
