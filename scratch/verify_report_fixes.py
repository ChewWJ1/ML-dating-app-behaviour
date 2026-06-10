"""
Verification script to audit report fixes in WIA1006 ML Report V8 (final).docx.
Checks that all 14 targeted edits are present and correct, and no old values remain.
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

DOCX_PATH = r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx'
doc = Document(DOCX_PATH)

print("=== STARTING PROGRAMMATIC AUDIT ===")

errors = []

# Helper to check text is present and old is not
def check_paragraph(p_idx, old_sub, new_sub):
    txt = doc.paragraphs[p_idx].text
    if old_sub in txt:
        errors.append(f"Para {p_idx} still contains old substring: '{old_sub}'")
    if new_sub not in txt:
        errors.append(f"Para {p_idx} missing new substring: '{new_sub}'")
        print(f"Actual text in Para {p_idx}: {txt}")

check_paragraph(27, "14 custom architectures and 2 AutoML baselines (total 16 models)", "11 custom architectures and 2 AutoML baselines (total 13 models)")
check_paragraph(163, "We evaluated 16 diverse models to establish performance baselines.", "We evaluated 13 diverse models to establish performance baselines.")
check_paragraph(183, "all 14 baseline and advanced models are tabulated below:", "all 13 baseline and advanced models are tabulated below:")
check_paragraph(395, "We evaluated 16 baseline and advanced classifiers", "We evaluated 13 baseline and advanced classifiers")
check_paragraph(395, "MLP, Label Smoothing MLP, Cosine KNN CF, GAT Graph Neural Network, SCARF Contrastive Learner, Opacus DP-SGD, TabPFN Zero-Shot, and TabNet-style Attentive Net", "Balanced Random Forest, KNN (Cosine Metric), FT-Transformer, SAINT, and NODE")

check_paragraph(211, "Figure 24: 5-Fold Cross-Validation", "Figure 24: 10-Fold Cross-Validation")
check_paragraph(212, "presents the 5-fold cross-validation accuracies", "presents the 10-fold cross-validation accuracies")
check_paragraph(240, "using a 5-fold CV RandomizedSearchCV.", "using a 10-fold CV RandomizedSearchCV.")
check_paragraph(367, "across 5 independent validation folds.", "across 10 independent validation folds.")
check_paragraph(367, "statistically weak 5-fold paired t-test", "statistically weak 10-fold paired t-test")
check_paragraph(379, "distribute the 5 validation folds cleanly", "distribute the 10 validation folds cleanly")

check_paragraph(411, "compressed student logistic regression model.", "compressed student lightweight neural network (StudentNet) model.")

# Check Tables
def check_table_cell(t_idx, r_idx, c_idx, old_sub, new_sub):
    txt = doc.tables[t_idx].rows[r_idx].cells[c_idx].text
    if old_sub in txt:
        errors.append(f"Table {t_idx}, R{r_idx}, C{c_idx} still contains old substring: '{old_sub}'")
    if new_sub not in txt:
        errors.append(f"Table {t_idx}, R{r_idx}, C{c_idx} missing new substring: '{new_sub}'")

check_table_cell(4, 6, 2, "5-fold CV", "10-fold CV")
check_table_cell(9, 2, 2, "5-fold cross-validation", "10-fold cross-validation")
check_table_cell(10, 2, 3, "5-fold cross-validation", "10-fold cross-validation")

print()
if errors:
    print(f"❌ AUDIT FAILED with {len(errors)} errors:")
    for err in errors:
        print(f"  - {err}")
    sys.exit(1)
else:
    print("✅ AUDIT PASSED! All 14 changes verified successfully.")
