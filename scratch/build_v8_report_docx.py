import os
import shutil
import zipfile
import tempfile
import re
from docx import Document

def replace_text_in_runs(paragraph, old_text, new_text):
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)

def update_document_xml(xml_content):
    # Safe text replacements in XML
    # Using regex to replace content within <w:t> tags
    # This prevents breaking the XML structure.
    
    # We decode to string
    content = xml_content.decode('utf-8')
    
    # Text replacements
    # Since word might split "V5.2" into multiple tags like <w:t>V5.</w:t> <w:t>2</w:t>, 
    # we can remove tags in between if we know they exist, but it's risky.
    # The safest way is to just do straight replacements if they happen to be in one <w:t> tag,
    # or just use docx for text replacement if XML fails.
    
    # Let's replace versions globally
    content = content.replace("V5.2", "V8")
    content = content.replace("V5.1", "V8")
    
    # Return encoded
    return content.encode('utf-8')

def patch_images_and_xml(original_docx, new_docx, mapping, plots_dir):
    print(f"Extracting {original_docx} to modify...")
    with tempfile.TemporaryDirectory() as tmpdir:
        with zipfile.ZipFile(original_docx, 'r') as zip_ref:
            zip_ref.extractall(tmpdir)
        
        # 1. Replace images
        for docx_image, plot_filename in mapping.items():
            source_path = os.path.join(plots_dir, plot_filename)
            target_path = os.path.join(tmpdir, docx_image)
            
            if os.path.exists(source_path):
                # We need to make sure we keep the same file extension or Word will complain.
                # Word's imageX.png must remain a .png
                shutil.copy(source_path, target_path)
                print(f"Successfully replaced {docx_image} with {plot_filename}")
            else:
                print(f"WARNING: Source image {source_path} not found!")
                
        # 2. Update document.xml
        xml_path = os.path.join(tmpdir, 'word', 'document.xml')
        with open(xml_path, 'rb') as f:
            xml_content = f.read()
            
        new_xml_content = update_document_xml(xml_content)
        
        with open(xml_path, 'wb') as f:
            f.write(new_xml_content)
                
        # 3. Zip back
        print(f"Creating new document {new_docx}...")
        def zipdir(path, ziph):
            for root, dirs, files in os.walk(path):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, path)
                    ziph.write(file_path, arcname)
                    
        with zipfile.ZipFile(new_docx, 'w', zipfile.ZIP_DEFLATED) as zipf:
            zipdir(tmpdir, zipf)
    print("Zip modification complete.")

def main():
    root_dir = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour"
    original_docx = os.path.join(root_dir, "reports", "WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.2 (final).docx")
    new_docx = os.path.join(root_dir, "reports", "WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx")
    plots_dir = os.path.join(root_dir, "assets", "v8 plots")

    # Image mapping
    mapping = {
        "word/media/image6.png": "01_match_outcome_distribution.png",
        "word/media/image7.png": "02_categorical_feature_distributions.png",
        "word/media/image8.png": "03_numerical_feature_distributions.png",
        "word/media/image9.png": "04_outlier_detection_boxplots.png",
        "word/media/image10.png": "05_numerical_features_by_match_outcome.png",
        "word/media/image11.png": "06_categorical_features_by_match_outcome.png",
        "word/media/image12.png": "07_pearson_correlation_matrix.png",
        "word/media/image13.png": "08_top_30_interest_tags.png",
        "word/media/image14.png": "09_causal_dag.png",
        "word/media/image15.png": "10_causal_adjacency_matrix.png",
        "word/media/image16.png": "11_ood_anomaly_score_distribution.png",
        "word/media/image17.png": "12_anova_f_scores.png",
        "word/media/image18.png": "13_mutual_information_scores.png",
        "word/media/image19.png": "14_pca_explained_variance.png",
        "word/media/image20.png": "15_pca_biplot.png",
        "word/media/image23.png": "16_train_test_class_distribution.png",
        "word/media/image24.png": "18_model_performance_comparison.png",
        "word/media/image25.png": "19_confusion_matrices.png",
        "word/media/image26.png": "20_roc_curves.png",
        "word/media/image27.png": "21_cross_validation_accuracy.png",
        "word/media/image28.png": "22_learning_curves.png",
        "word/media/image29.png": "36_probability_calibration_reliability_diagram.png",
        "word/media/image30.png": "26_baseline_vs_tuned_comparison.png",
        "word/media/image31.png": "27_best_tuned_model_details.png",
        "word/media/image32.png": "28_feature_importance.png",
        "word/media/image33.png": "30_friedman_h_statistic_heatmap.png",
        "word/media/image34.png": "31_shap_dependence_interaction_plot.png",
        "word/media/image35.png": "32_shap_interaction_matrix_heatmap.png",
        "word/media/image36.png": "33_conformal_prediction_set_sizes.png",
        "word/media/image37.png": "34_bayesian_uncertainty_mc_dropout.png",
        "word/media/image38.png": "35_adversarial_robustness_fgsm.png",
        "word/media/image39.png": "38_causal_uplift_targeting_segments.png",
        "word/media/image41.png": "37_knowledge_distillation_teacher_student_comparison.png",
        "word/media/image42.png": "24_attentive_tabular_network_feature_selection.png",
        "word/media/image43.png": "25_scarf_contrastive_learning_embeddings.png",
        "word/media/image44.png": "17_label_smoothing_mixup_regularization.png",
        "word/media/image45.png": "23_differential_privacy_comparison.png",
    }

    # Step 1: Zip-level modification for images and raw XML string replacement
    patch_images_and_xml(original_docx, new_docx, mapping, plots_dir)
    
    # Step 2: Use python-docx to update specific table metrics safely
    print("Updating tables via python-docx...")
    doc = Document(new_docx)
    
    # Replace some specific full-paragraph strings to handle split words in XML
    for p in doc.paragraphs:
        if "V5.2" in p.text:
            p.text = p.text.replace("V5.2", "V8")
        if "V5.1" in p.text:
            p.text = p.text.replace("V5.1", "V8")
        if "V6" in p.text:
            # Dangerous if there are other "V6"s, so we'll be careful
            # Let's only replace it if it's " V6"
            if " V6" in p.text:
                p.text = p.text.replace(" V6", " V9")

    # In Table 6 (Baseline) or Table 7 (Tuned), we update the champion details
    # We won't programmatically rewrite the whole table layout to avoid corruption,
    # but we can replace specific text in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Update text inside table cells if needed
                if "Random Forest" in cell.text and "Champion" in cell.text:
                    cell.text = cell.text.replace("Random Forest", "LightGBM (Tuned)")
                if "0.5143" in cell.text: # Old SVM ROC AUC
                    pass # Baseline SVM is still 0.5143 according to plan
                
    doc.save(new_docx)
    print("Done! Generated V8 report.")

if __name__ == "__main__":
    main()
