import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.2 (final).docx"
doc = docx.Document(doc_path)

print("--- Searching for DML or 3.2.1 ---")
for idx, para in enumerate(doc.paragraphs):
    txt = para.text
    if "Double Machine Learning" in txt or "DML" in txt or "3.2.1" in txt:
        print(f"P{idx}: '{txt[:150]}'")
        # print surrounding
        for i in range(max(0, idx-2), min(len(doc.paragraphs), idx+3)):
            if i != idx:
                print(f"  P{i}: '{doc.paragraphs[i].text.strip()}'")
