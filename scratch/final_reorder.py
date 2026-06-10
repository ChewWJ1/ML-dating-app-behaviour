"""
Final precise reorder — fix remaining out-of-order figure blocks.
We now directly address each figure block by its anchor paragraph,
then deterministically extract and reinsert in exact order.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from lxml import etree

DOCX_PATH = r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx'
doc = Document(DOCX_PATH)

def has_image(p):
    xml = etree.tostring(p._element, encoding='unicode')
    return '<w:drawing' in xml or '<v:shape' in xml

def show_range(paras, a, b):
    for i in range(a, min(b, len(paras))):
        p = paras[i]
        img = '[IMG]' if has_image(p) else '     '
        print(f"  {i:4d} {img} {p.text[:105]}")

def reinsert(paras, indices_in_correct_order, anchor_text):
    """Remove elements at given para indices, reinsert in order before anchor."""
    # Find anchor
    anchor_elem = None
    for p in paras:
        if anchor_text in p.text:
            anchor_elem = p._element
            break
    if not anchor_elem:
        print(f"  ERROR: anchor not found: {anchor_text[:50]}")
        return False
    elems = [paras[i]._element for i in indices_in_correct_order]
    for e in elems:
        e.getparent().remove(e)
    for e in reversed(elems):
        anchor_elem.addprevious(e)
    return True

paras = doc.paragraphs
print("=== PRE-FIX STATE ===")

# ── FIG 41 ──────────────────────────────────────────────────────────────────
# Look for the block
print("\nFig 41 block:")
show_range(paras, 233, 242)
# We see: body(235), caption(237), intro(238) — wrong. Correct: intro, caption, [IMG], body
# Identify exact indices
f41 = {}
for i, p in enumerate(paras):
    if "Figure 41 below presents" in p.text:       f41['intro']   = i
    elif "Figure 41: PyTorch BCE Loss" in p.text:  f41['caption'] = i
    elif "Figure 41 compares the training loss curves of standard Binary" in p.text: f41['body'] = i
for i, p in enumerate(paras):
    if has_image(p) and any(abs(i - v) <= 4 for v in f41.values()) and 'image' not in f41:
        # Verify it's not one of the text paras
        if p.text.strip() == '':
            f41['image'] = i; break

print("  f41:", f41)
# Correct order: intro, caption, image, body
if len(f41) >= 3:
    order_41 = []
    if 'intro'   in f41: order_41.append(f41['intro'])
    if 'caption' in f41: order_41.append(f41['caption'])
    if 'image'   in f41: order_41.append(f41['image'])
    if 'body'    in f41: order_41.append(f41['body'])
    reinsert(paras, order_41, "4.3 Hyperparameter Tuning")
    print("  Fig 41 reinserted in order:", order_41)

doc.save(DOCX_PATH)
doc = Document(DOCX_PATH); paras = doc.paragraphs

# ── FIG 42 ──────────────────────────────────────────────────────────────────
print("\nFig 42 block:")
f42 = {}
for i, p in enumerate(paras):
    if "Figure 42 below tracks the Opacus" in p.text:    f42['intro']   = i
    elif "Figure 42: Opacus Differential Privacy" in p.text: f42['caption'] = i
    elif "Figure 42 tracks the privacy budget consumption" in p.text: f42['body'] = i
for i, p in enumerate(paras):
    if has_image(p) and any(abs(i - v) <= 4 for v in f42.values()) and 'image' not in f42:
        if p.text.strip() == '':
            f42['image'] = i; break

show_range(paras, min(f42.values()) - 1 if f42 else 263, max(f42.values()) + 2 if f42 else 268)
print("  f42:", f42)

if len(f42) >= 2:
    order_42 = []
    if 'intro'   in f42: order_42.append(f42['intro'])
    if 'caption' in f42: order_42.append(f42['caption'])
    if 'image'   in f42: order_42.append(f42['image'])
    if 'body'    in f42: order_42.append(f42['body'])
    reinsert(paras, order_42, "5.0 Insights and Interpretation")
    print("  Fig 42 reinserted in order:", order_42)

doc.save(DOCX_PATH)
doc = Document(DOCX_PATH); paras = doc.paragraphs

# ── FIGs 39+40 ──────────────────────────────────────────────────────────────
print("\nFigs 39+40 block:")
intro_idx = None; f39 = {}; f40 = {}
for i, p in enumerate(paras):
    if "Complementing the SHAP attributions above" in p.text: intro_idx = i
    elif "Figure 39: Attentive TabNet" in p.text:  f39['caption'] = i
    elif "attentive feature selection heatmap in Figure 39" in p.text: f39['body'] = i
    elif "Figure 40: SCARF Self-Supervised" in p.text: f40['caption'] = i
    elif "Figure 40 projects the SCARF contrastive" in p.text: f40['body'] = i
    elif "Methodological Disclosure: During our SCARF" in p.text: f40['disclosure'] = i

# Find images: the image for Fig39 should be between Fig39 caption and body
# The image for Fig40 should be between Fig40 caption and body
all_img_indices = [i for i, p in enumerate(paras) if has_image(p) and p.text.strip() == '']
# Filter to those in the relevant range
relevant_range_start = intro_idx if intro_idx else 285
relevant_range_end   = (f40.get('disclosure', f39.get('body', relevant_range_start)) or relevant_range_start) + 5
block_imgs = [i for i in all_img_indices if relevant_range_start - 5 <= i <= relevant_range_end]
print(f"  Image candidates in block: {block_imgs}")
# Assign: if 2 images, first goes to fig39, second to fig40
if len(block_imgs) >= 2:
    f39['image'] = block_imgs[0]
    f40['image'] = block_imgs[1]
elif len(block_imgs) == 1:
    # Assign to whichever figure has caption closest
    f39['image'] = block_imgs[0]

print("  intro:", intro_idx, "f39:", f39, "f40:", f40)
show_range(paras, (intro_idx or 288) - 1, max(
    [v for v in list(f39.values()) + list(f40.values()) + [intro_idx or 0] if v] + [0]
) + 3)

# Correct order: intro, Fig39_caption, Fig39_image, Fig39_body, Fig40_caption, Fig40_image, Fig40_body, disclosure
order_3940 = []
if intro_idx is not None:         order_3940.append(intro_idx)
if 'caption' in f39:              order_3940.append(f39['caption'])
if 'image' in f39:                order_3940.append(f39['image'])
if 'body' in f39:                 order_3940.append(f39['body'])
if 'caption' in f40:              order_3940.append(f40['caption'])
if 'image' in f40:                order_3940.append(f40['image'])
if 'body' in f40:                 order_3940.append(f40['body'])
if 'disclosure' in f40:           order_3940.append(f40['disclosure'])

if order_3940:
    reinsert(paras, order_3940, "5.3 Demographic Parity")
    print("  Figs 39+40 reinserted in order:", order_3940)

doc.save(DOCX_PATH)
doc = Document(DOCX_PATH); paras = doc.paragraphs

# ── FIG 38 ──────────────────────────────────────────────────────────────────
print("\nFig 38 block:")
intro_38 = None; f38 = {}
for i, p in enumerate(paras):
    if "Figure 38 below presents the training loss curves comparing" in p.text: intro_38 = i
    elif "Figure 38: Knowledge Distillation Student" in p.text: f38['caption'] = i
    elif "Figure 38 compares the training loss profiles of the teacher ensemble" in p.text: f38['body'] = i
all_img = [i for i, p in enumerate(paras) if has_image(p) and p.text.strip() == '']
cap38 = f38.get('caption')
if cap38:
    close_imgs = [i for i in all_img if abs(i - cap38) <= 5]
    if close_imgs: f38['image'] = close_imgs[0]

print("  intro:", intro_38, "f38:", f38)
show_range(paras, (intro_38 or 385) - 1, max([v for v in list(f38.values()) + [intro_38 or 0] if v] + [0]) + 3)

order_38 = []
if intro_38 is not None:  order_38.append(intro_38)
if 'caption' in f38:      order_38.append(f38['caption'])
if 'image' in f38:        order_38.append(f38['image'])
if 'body' in f38:         order_38.append(f38['body'])

if order_38:
    reinsert(paras, order_38, "7.3 Summary of Evaluated and Excluded")
    print("  Fig 38 reinserted in order:", order_38)

doc.save(DOCX_PATH)
print("\nDocument saved. All done.")

# ── FINAL VERIFICATION ──────────────────────────────────────────────────────
doc = Document(DOCX_PATH); paras = doc.paragraphs
print("\n=== FINAL VERIFICATION ===")

print("\nFig 41 final:")
for i, p in enumerate(paras):
    if any(kw in p.text for kw in ["Figure 41", "Mixup", "BCE", "Label-Smoothed BCE", "4.3 Hyper"]):
        img = '[IMG]' if has_image(p) else ''
        print(f"  {i:4d} {img:5} {p.text[:100]}")

print("\nFig 42 final:")
for i, p in enumerate(paras):
    if any(kw in p.text for kw in ["Figure 42", "Opacus", "DP-SGD", "epsilon", "5.0 Insights"]):
        img = '[IMG]' if has_image(p) else ''
        print(f"  {i:4d} {img:5} {p.text[:100]}")

print("\nFigs 39+40 final:")
for i, p in enumerate(paras):
    if any(kw in p.text for kw in ["Figure 39", "Figure 40", "SCARF", "AttentiveTab", "5.3 Demo"]):
        img = '[IMG]' if has_image(p) else ''
        print(f"  {i:4d} {img:5} {p.text[:100]}")

print("\nFig 38 final:")
for i, p in enumerate(paras):
    if any(kw in p.text for kw in ["Figure 38", "Knowledge Distill", "StudentNet", "7.3 Summary"]):
        img = '[IMG]' if has_image(p) else ''
        print(f"  {i:4d} {img:5} {p.text[:100]}")

print("\nSection 8 check (should be clean):")
in_sec8 = False
for i, p in enumerate(paras):
    if "8.1 Key Findings" in p.text or "8.2 Recommendations" in p.text: in_sec8 = True
    if "9.0 References" in p.text: in_sec8 = False
    if in_sec8:
        img = '[IMG]' if has_image(p) else ''
        if p.text.strip() or img:
            print(f"  {i:4d} {img:5} {p.text[:100]}")
