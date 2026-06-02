import docx
import sys
import os

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.2 (final).docx"
doc = docx.Document(doc_path)

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")

print("\n--- Listing Tables ---")
for idx, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns) if rows > 0 else 0
    print(f"Table {idx}: Rows={rows}, Cols={cols}")
    if rows > 0:
        # print first row text (header)
        header_text = [cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells]
        print(f"  Header: {header_text[:4]}...")

print("\n--- Search for Figures or Captions ---")
for idx, para in enumerate(doc.paragraphs):
    txt = para.text.strip()
    if "Figure " in txt or "Table " in txt:
        print(f"P{idx}: {txt[:120]}")
