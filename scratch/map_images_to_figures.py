import docx
import sys
import os

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.2 (final).docx"
doc = docx.Document(doc_path)

# Build a mapping of relationship ID to image target path
rels = doc.part.rels
rId_to_image = {}
for rId, rel in rels.items():
    if "image" in rel.target_ref:
        rId_to_image[rId] = rel.target_ref
        print(f"Rel: {rId} -> {rel.target_ref}")

print("\n--- Mapping images in paragraphs ---")
for p_idx, para in enumerate(doc.paragraphs):
    # Search for drawings/images in the paragraph XML
    import re
    xml_str = para._p.xml
    # Find all rId references in the XML
    rIds = re.findall(r'r:embed="([^"]+)"', xml_str)
    rIds.extend(re.findall(r'r:link="([^"]+)"', xml_str))
    
    if rIds:
        unique_rIds = list(set(rIds))
        resolved = [rId_to_image.get(rId, f"Unknown ({rId})") for rId in unique_rIds if rId in rId_to_image]
        if resolved:
            print(f"P{p_idx}: Text='{para.text.strip()}' -> Images: {resolved}")
            # print surrounding paragraphs
            start = max(0, p_idx - 2)
            end = min(len(doc.paragraphs), p_idx + 3)
            for i in range(start, end):
                if i != p_idx:
                    print(f"  P{i}: '{doc.paragraphs[i].text.strip()}'")

print("\n--- Mapping images in tables ---")
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            for p_idx, para in enumerate(cell.paragraphs):
                xml_str = para._p.xml
                import re
                rIds = re.findall(r'r:embed="([^"]+)"', xml_str)
                rIds.extend(re.findall(r'r:link="([^"]+)"', xml_str))
                if rIds:
                    unique_rIds = list(set(rIds))
                    resolved = [rId_to_image.get(rId, f"Unknown ({rId})") for rId in unique_rIds if rId in rId_to_image]
                    if resolved:
                        print(f"Table {t_idx}, Row {r_idx}, Col {c_idx}, P{p_idx}: Text='{para.text.strip()}' -> Images: {resolved}")
