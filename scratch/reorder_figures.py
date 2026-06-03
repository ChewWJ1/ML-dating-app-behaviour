"""
Fix the ordering of the moved figure blocks. The reversed() insertion in move_figures.py
put things in the wrong order. This script reorders each figure block in-place.

Current wrong order at each location:
  Fig 41 (paras 235-238): [IMG], body(236), caption(237), intro(238)
  Fig 42 (paras 264-266): body(264), caption(265), intro(266)
  Figs 39+40 (paras 290-297): [IMG](290), disclosure(291), body40(292), cap40(293),
                                [IMG](294), body39(295), cap39(296), intro(297)
  Fig 38 (paras 387-390):  [IMG](387), body(388), caption(389), intro(390)

Target correct order:
  Fig 41: intro, caption, [IMG], body
  Fig 42: intro, caption, [IMG], body
  Figs 39+40: intro, cap39, [IMG]39, body39, cap40, [IMG]40, body40, disclosure40
  Fig 38: intro, caption, [IMG], body
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from lxml import etree

DOCX_PATH = r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx'
doc = Document(DOCX_PATH)
body = doc.element.body

def has_image(para):
    xml = etree.tostring(para._element, encoding='unicode')
    return '<w:drawing' in xml or '<v:shape' in xml

def reorder_elems(elem_list, first_elem):
    """
    Remove all elems in elem_list from body, then re-insert them in order before first_elem.
    first_elem is the para that comes AFTER the block (the anchor).
    """
    # Detach all
    parent = first_elem.getparent()
    for e in elem_list:
        e.getparent().remove(e)
    # Re-insert in order before the anchor
    for e in reversed(elem_list):
        first_elem.addprevious(e)

# ─────────────────────────────────────────────────────────────────────────────
# Reload current state to get fresh para indices
print("Analyzing current paragraph structure...")
paras = doc.paragraphs

def dump_range(start, end):
    for i in range(start, end):
        p = paras[i]
        img = '[IMG]' if has_image(p) else '     '
        print(f"  {i:4d} {img} {p.text[:100]}")

# ─────────────────────────────────────────────────────────────────────────────
# FIG 41: Find the block around "Figure 41"
print("\n=== Fixing Fig 41 ===")
# Find the intro sentence (contains "Label-Smoothed Mixup" and is after para 230ish)
f41_idx = {}
for i, p in enumerate(paras):
    t = p.text
    if "Figure 41 below presents the training loss comparison" in t:
        f41_idx['intro'] = i
    elif "Figure 41: PyTorch BCE Loss" in t:
        f41_idx['caption'] = i
    elif "Figure 41 compares the training loss curves of standard Binary" in t:
        f41_idx['body'] = i
    if has_image(p) and i > 230 and i < 245 and 'image' not in f41_idx:
        # Check it's the right image - near the other fig41 paras
        f41_idx['image_candidate'] = i

print("  Found:", f41_idx)
dump_range(233, 241)

# Correct order: intro -> caption -> image -> body
# Find the image para between intro and 4.3 heading
img_para = None
for i in range(233, 241):
    if has_image(paras[i]) and i not in [f41_idx.get('intro'), f41_idx.get('caption'), f41_idx.get('body')]:
        img_para = i
        break

if 'intro' in f41_idx and 'caption' in f41_idx and 'body' in f41_idx:
    intro_e   = paras[f41_idx['intro']]._element
    caption_e = paras[f41_idx['caption']]._element
    body_e    = paras[f41_idx['body']]._element
    img_e     = paras[img_para]._element if img_para else None

    # Find anchor = "4.3 Hyperparameter" paragraph
    anchor_4_3 = None
    for i, p in enumerate(paras):
        if "4.3 Hyperparameter Tuning" in p.text:
            anchor_4_3 = p._element
            print(f"  Anchor 4.3 found at para {i}")
            break

    if anchor_4_3:
        # Remove all block elements
        for e in [intro_e, caption_e, body_e] + ([img_e] if img_e else []):
            e.getparent().remove(e)
        # Also remove the blank para that was between fig26 and the block
        # Re-insert in correct order
        ordered = [intro_e, caption_e] + ([img_e] if img_e else []) + [body_e]
        for e in reversed(ordered):
            anchor_4_3.addprevious(e)
        print("  Fig 41 reordered: intro -> caption -> image -> body")

# ─────────────────────────────────────────────────────────────────────────────
# FIG 42: Find the block
# Must reload since we modified structure
doc.save(DOCX_PATH)
doc = Document(DOCX_PATH)
paras = doc.paragraphs

print("\n=== Fixing Fig 42 ===")
f42_idx = {}
for i, p in enumerate(paras):
    t = p.text
    if "Figure 42 below tracks the Opacus DP-SGD privacy budget" in t:
        f42_idx['intro'] = i
    elif "Figure 42: Opacus Differential Privacy" in t:
        f42_idx['caption'] = i
    elif "Figure 42 tracks the privacy budget consumption (epsilon)" in t:
        f42_idx['body'] = i

print("  Found:", f42_idx)
# Check if there's an image in fig42 block area
for i in range(f42_idx.get('body', 260)-2, f42_idx.get('intro', 270)+2):
    if i < len(paras) and has_image(paras[i]):
        f42_idx['image'] = i
        break

print("  Block area:")
start = min(f42_idx.values()) - 1 if f42_idx else 263
dump_range(start, start+6)

# Find anchor = "5.0 Insights"
anchor_5_0 = None
for i, p in enumerate(paras):
    if "5.0 Insights and Interpretation" in p.text:
        anchor_5_0 = p._element
        print(f"  Anchor 5.0 found at para {i}")
        break

if anchor_5_0 and 'intro' in f42_idx and 'caption' in f42_idx and 'body' in f42_idx:
    intro_e   = paras[f42_idx['intro']]._element
    caption_e = paras[f42_idx['caption']]._element
    body_e    = paras[f42_idx['body']]._element
    img_e     = paras[f42_idx['image']]._element if 'image' in f42_idx else None

    for e in [intro_e, caption_e, body_e] + ([img_e] if img_e else []):
        e.getparent().remove(e)

    ordered = [intro_e, caption_e] + ([img_e] if img_e else []) + [body_e]
    for e in reversed(ordered):
        anchor_5_0.addprevious(e)
    print("  Fig 42 reordered: intro -> caption -> image -> body")

# ─────────────────────────────────────────────────────────────────────────────
# FIGs 39+40: Correct order should be:
# intro_text, Fig39_caption, Fig39_image, Fig39_body, Fig40_caption, Fig40_image, Fig40_body, Fig40_disclosure
doc.save(DOCX_PATH)
doc = Document(DOCX_PATH)
paras = doc.paragraphs

print("\n=== Fixing Figs 39+40 ===")
f39 = {}; f40 = {}; intro_3940 = None
for i, p in enumerate(paras):
    t = p.text
    if "Complementing the SHAP attributions above" in t:
        intro_3940 = i
    elif "Figure 39: Attentive TabNet" in t:
        f39['caption'] = i
    elif "attentive feature selection heatmap in Figure 39 projects" in t:
        f39['body'] = i
    elif "Figure 40: SCARF Self-Supervised" in t:
        f40['caption'] = i
    elif "Figure 40 projects the SCARF contrastive pre-trained" in t:
        f40['body'] = i
    elif "Methodological Disclosure: During our SCARF" in t:
        f40['disclosure'] = i

# Find images by proximity to their captions
for i, p in enumerate(paras):
    if has_image(p):
        # Check if near fig39 caption
        if 'caption' in f39 and abs(i - f39['caption']) <= 3 and 'image' not in f39:
            f39['image'] = i
        elif 'caption' in f40 and abs(i - f40['caption']) <= 4 and 'image' not in f40:
            f40['image'] = i

print("  Fig39:", f39)
print("  Fig40:", f40)
print("  intro:", intro_3940)

dump_range(min(intro_3940 or 289, min([v for v in list(f39.values())+list(f40.values()) if isinstance(v,int)]))-1,
           max([v for v in list(f39.values())+list(f40.values()) if isinstance(v,int)])+3)

# Find anchor = "5.3 Demographic Parity"
anchor_5_3 = None
for i, p in enumerate(paras):
    if "5.3 Demographic Parity" in p.text:
        anchor_5_3 = p._element
        print(f"  Anchor 5.3 at para {i}")
        break

if anchor_5_3:
    elems = []
    if intro_3940 is not None: elems.append(paras[intro_3940]._element)
    # Fig 39 block: caption, image, body
    if 'caption' in f39: elems.append(paras[f39['caption']]._element)
    if 'image' in f39:   elems.append(paras[f39['image']]._element)
    if 'body' in f39:    elems.append(paras[f39['body']]._element)
    # Fig 40 block: caption, image, body, disclosure
    if 'caption' in f40: elems.append(paras[f40['caption']]._element)
    if 'image' in f40:   elems.append(paras[f40['image']]._element)
    if 'body' in f40:    elems.append(paras[f40['body']]._element)
    if 'disclosure' in f40: elems.append(paras[f40['disclosure']]._element)

    for e in elems:
        e.getparent().remove(e)
    for e in reversed(elems):
        anchor_5_3.addprevious(e)
    print("  Figs 39+40 reordered correctly.")

# ─────────────────────────────────────────────────────────────────────────────
# FIG 38
doc.save(DOCX_PATH)
doc = Document(DOCX_PATH)
paras = doc.paragraphs

print("\n=== Fixing Fig 38 ===")
f38 = {}; intro_38 = None
for i, p in enumerate(paras):
    t = p.text
    if "Figure 38 below presents the training loss curves comparing" in t:
        intro_38 = i
    elif "Figure 38: Knowledge Distillation Student" in t:
        f38['caption'] = i
    elif "Figure 38 compares the training loss profiles of the teacher ensemble" in t:
        f38['body'] = i

for i, p in enumerate(paras):
    if has_image(p) and 'caption' in f38 and abs(i - f38['caption']) <= 4 and 'image' not in f38:
        f38['image'] = i; break

print("  Fig38:", f38, "intro:", intro_38)
if intro_38:
    dump_range(intro_38 - 1, (max([v for v in list(f38.values()) if isinstance(v,int)]) if f38 else intro_38) + 3)

anchor_7_3 = None
for i, p in enumerate(paras):
    if "7.3 Summary of Evaluated and Excluded" in p.text:
        anchor_7_3 = p._element
        print(f"  Anchor 7.3 at para {i}")
        break

if anchor_7_3 and intro_38 is not None and 'caption' in f38:
    elems = [paras[intro_38]._element]
    if 'caption' in f38: elems.append(paras[f38['caption']]._element)
    if 'image' in f38:   elems.append(paras[f38['image']]._element)
    if 'body' in f38:    elems.append(paras[f38['body']]._element)

    for e in elems:
        e.getparent().remove(e)
    for e in reversed(elems):
        anchor_7_3.addprevious(e)
    print("  Fig 38 reordered correctly.")

# ─────────────────────────────────────────────────────────────────────────────
doc.save(DOCX_PATH)
print("\nAll reordering complete. Document saved.")

# Final verification
doc = Document(DOCX_PATH)
paras = doc.paragraphs

print("\n=== FINAL CHECK ===")
print("-- Fig 41 area --")
for i, p in enumerate(paras):
    if any(kw in p.text for kw in ["Figure 41", "Mixup", "Label-Smoothed", "4.3 Hyperparameter"]):
        img = '[IMG]' if has_image(p) else '     '
        print(f"  {i:4d} {img} {p.text[:100]}")

print("\n-- Fig 42 area --")
for i, p in enumerate(paras):
    if any(kw in p.text for kw in ["Figure 42", "Opacus", "DP-SGD", "5.0 Insights"]):
        img = '[IMG]' if has_image(p) else '     '
        print(f"  {i:4d} {img} {p.text[:100]}")

print("\n-- Figs 39+40 area --")
for i, p in enumerate(paras):
    if any(kw in p.text for kw in ["Figure 39", "Figure 40", "SCARF", "AttentiveTab", "5.3 Demographic"]):
        img = '[IMG]' if has_image(p) else '     '
        print(f"  {i:4d} {img} {p.text[:100]}")

print("\n-- Fig 38 area --")
for i, p in enumerate(paras):
    if any(kw in p.text for kw in ["Figure 38", "Knowledge Distill", "StudentNet", "7.3 Summary"]):
        img = '[IMG]' if has_image(p) else '     '
        print(f"  {i:4d} {img} {p.text[:100]}")

print("\n-- Section 8.2 (should have NO figure paragraphs) --")
in_8 = False
for i, p in enumerate(paras):
    if "8.2 Recommendations" in p.text: in_8 = True
    if "9.0 References" in p.text: in_8 = False
    if in_8 and any(f"Figure {n}" in p.text for n in [38, 39, 40, 41, 42]):
        print(f"  STILL IN 8.2: {i:4d} {p.text[:100]}")
print("  8.2 check done.")
