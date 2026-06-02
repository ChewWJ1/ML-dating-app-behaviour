import docx
import os

root_dir = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour"

def inspect_file(name):
    path = os.path.join(root_dir, "reports", name)
    print(f"\nDocument: {name}")
    print(f"Path: {path}")
    if not os.path.exists(path):
        print("Does not exist!")
        return
    print(f"Size: {os.path.getsize(path):,} bytes")
    try:
        doc = docx.Document(path)
        print(f"Paragraphs: {len(doc.paragraphs)}")
        print(f"Tables: {len(doc.tables)}")
        
        # Print headings
        headings = []
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text.startswith(("1.0", "2.0", "3.0", "4.0", "5.0", "6.0", "7.0", "8.0", "9.0", "10.0")):
                headings.append((i, text))
        print(f"Headings count: {len(headings)}")
        for i, h in headings[:15]:
            print(f"  P{i}: {h}")
    except Exception as e:
        print(f"Error reading docx: {e}")

inspect_file("WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx")
inspect_file("WIA1006_Assignment_Report_V8_Patched.docx")
inspect_file("WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.2 (final).docx")
