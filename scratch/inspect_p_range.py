import docx
import os

doc_path = r"reports/WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.1(long).docx"
if not os.path.exists(doc_path):
    print("Error: File not found!")
    exit(1)

doc = docx.Document(doc_path)
print(f"Total paragraphs: {len(doc.paragraphs)}")

for idx in range(285, 316):
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        print(f"P {idx} ({p.style.name}): {repr(p.text)}")
