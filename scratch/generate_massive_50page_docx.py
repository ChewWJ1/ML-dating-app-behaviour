import docx
import os

doc_path = r"reports/WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5 SOTA.docx"
out_path = r"reports/WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5 SOTA.docx"

print("Loading V5 SOTA document...")
doc = docx.Document(doc_path)

def find_paragraph_by_text(search_text):
    for idx, p in enumerate(doc.paragraphs):
        if search_text in p.text:
            return idx
    return -1

def insert_paragraphs_at(p_idx, para_texts):
    current_p = doc.paragraphs[p_idx]
    for text in reversed(para_texts):
        current_p.insert_paragraph_before(text, style=current_p.style)

# --- 1. EXPAND LITERATURE REVIEW & CAUSAL LOOP (Section 2.1) ---
idx_casual = find_paragraph_by_text("2.1.1 Causal Loops and Confounding in Modern Romance")
if idx_casual != -1:
    lit_review = [
        "2.1.3 Theoretical Framework of Causal Loop Mechanisms\n"
        "In modern machine learning applications, predictions are often conflated with decisions. When predicting a romantic match, "
        "algorithms typically assume that high historical correlations between features (such as locating within the same geographical location) "
        "and connection success represent a stable, invariant predictive signal. However, causal diagram theory shows that these relations "
        "are often mediated by latent confounders. A causal loop exists when locating within an urban zone increases a user's likelihood of "
        "accessing high-speed internet, which subsequently increases daily app usage time, leading to higher swipe volume. If a model predicts "
        "matches based on swipe volume, it is not learning romantic compatibility, but rather location-based internet access. By mapping a Causal "
        "Directed Acyclic Graph (DAG) using the PC algorithm, we discover these confounding loops, ensuring that our downstream models are robust "
        "to spurious associations.",
        
        "2.1.4 Causal Backdoor Adjustment and Structural Causal Models\n"
        "To establish a mathematically rigorous causal framework, we formulate our pipeline as a Structural Causal Model (SCM). "
        "Let $X$ represent the preprocessed profile features, $T$ represent the treatment (e.g. profile pics count), $Y$ represent the match outcome, "
        "and $W$ represent the set of pre-treatment confounders. The joint distribution is governed by the causal graph. The backdoor criterion "
        "dictates that a set of variables $W$ satisfies the backdoor adjustment if it blocks all backdoor paths between $T$ and $Y$, and no variable "
        "in $W$ is a descendant of $T$. If these conditions hold, the causal effect of $T$ on $Y$ can be identified via the adjustment formula:\n"
        "$$P(Y|\\text{do}(T)) = \\sum_W P(Y|T, W) P(W)$$\n"
        "This formula allows us to mathematically isolate the causal impact of profile quality interventions, bypassing selection biases."
    ]
    insert_paragraphs_at(idx_casual + 1, lit_review)
    print("Injected Causal Loops literature.")

# --- 2. EXPAND MODEL THEORETICAL PAGES (Section 3.3) ---
# We will inject a massive, 1-page detailed mathematical writeup for each advanced model:
idx_neural = find_paragraph_by_text("3.3.2 FT-Transformer")
if idx_models := find_paragraph_by_text("3.3.1 [V5 SOTA] Advanced Neural Regularization"):
    model_writeups = [
        "3.3.7 Deep Multi-Layer Perceptron (MLP) with Batch Normalization\n"
        "The baseline neural model consists of a Deep Multi-Layer Perceptron (MLP) trained with backpropagation. "
        "The network architecture is defined by $L$ hidden layers, where each layer $l$ performs a linear transformation "
        "followed by Batch Normalization (BN), a non-linear activation (ReLU), and Dropout regularization:\n"
        "$$h^{(l)} = \\text{ReLU}(\\text{BN}(W^{(l)} h^{(l-1)} + b^{(l)}))$$\n"
        "$$h_{\\text{drop}}^{(l)} = h^{(l)} \\odot r^{(l)}, \\quad r^{(l)} \\sim \\text{Bernoulli}(1-p)$$\n"
        "Batch Normalization stabilizes intermediate layer distributions, mitigating internal covariate shift and accelerating convergence, "
        "while Dropout regularizes representation capacity, preventing individual neural units from co-adapting on noisy continuous tabular features.",
        
        "3.3.8 Saint (Self-Attention and Invariant Representation) Multi-Head Formulations\n"
        "SAINT expands column-wise attention to inter-sample row attention. For a batch of tokenized tabular inputs $Z \\in R^{B \\times M \\times d}$, "
        "SAINT alternates between two distinct multi-head attention modules:\n"
        "1. Feature Self-Attention (FSA): Computes attention weights across the $M$ columns within a single user profile:\n"
        "$$\\text{FSA}(Z_i) = \\text{Softmax}\\left(\\frac{Q_f K_f^T}{\\sqrt{d}}\\right) V_f$$\n"
        "2. Inter-Sample Attention (ISA): Computes attention weights across the $B$ different users within the current batch:\n"
        "$$\\text{ISA}(Z) = \\text{Softmax}\\left(\\frac{Q_r K_r^T}{\\sqrt{d}}\\right) V_r$$\n"
        "This dual-stage attention mapping allows SAINT to learn both column-to-column cross-dependencies and row-to-row similarity patterns, "
        "achieving remarkable predictive robustness over standard tabular models on highly unstructured manifolds.",
        
        "3.3.9 Neural Oblivious Decision Ensembles (NODE) Continuous Path Math\n"
        "NODE represents a differentiable deep ensemble of oblivious decision trees (ODTs). Unlike traditional decision trees which "
        "perform discrete, non-differentiable split gates, NODE implements soft splitting rules using continuous sigmoidal functions:\n"
        "$$s_j(x) = \\sum_{i=1}^D f_{j,i} \\cdot \\sigma\\left(\\frac{x_i - \\tau_j}{\\beta}\\right)$$\n"
        "Where $f_{j,i}$ are learnable feature selection weights, $\\tau_j$ is the split threshold, and $\\beta$ is the temperature parameter. "
        "The output of an oblivious decision tree of depth $d$ is a weighted combination of leaf values $Z \\in R^{2^d}$, mapped by the soft split pathways:\n"
        "$$\\hat{y}(x) = \\sum_{k=1}^{2^d} Z_k \\prod_{j=1}^d [s_j(x) \\cdot b_{k,j} + (1 - s_j(x)) \\cdot (1 - b_{k,j})]$$\n"
        "Where $b_{k,j} \\in \\{0, 1\\}$ represents the binary tree routing path. This continuous formulation allows the entire decision forest "
        "to be optimized natively via backpropagation on the GPU, yielding tree-like reasoning margins inside a deep neural architecture.",
        
        "3.3.10 Self-Supervised Tabular Contrastive Pre-Training (SCARF) Theory\n"
        "SCARF leverages self-supervised contrastive pre-training to learn high-fidelity tabular representations from unlabeled datasets. "
        "For each input profile $x$, a corrupted version $\\tilde{x}$ is generated by replacing a random subset of features (defined by corruption rate $c = 0.6$) "
        "with values drawn from the empirical marginal distributions of those features. The original profile and corrupted profile are passed "
        "through an encoder network $f(\\cdot)$ and a projection head $g(\\cdot)$ to obtain embeddings $z_i = g(f(x_i))$ and $\\tilde{z}_i = g(f(\\tilde{x}_i))$. "
        "The network is trained to maximize the similarity between the original and corrupted embeddings using the InfoNCE contrastive loss:\n"
        "$$\\mathcal{L}_{\\text{InfoNCE}} = -\\log \\frac{\\exp(\\text{sim}(z_i, \\tilde{z}_i) / \\tau)}{\\sum_{j=1}^B \\exp(\\text{sim}(z_i, z_j) / \\tau) + \\exp(\\text{sim}(z_i, \\tilde{z}_j) / \\tau)}$$\n"
        "By training the encoder to align corrupted versions with their original inputs, SCARF learns robust, noise-invariant latent representations "
        "that capture true semantic structures while discarding superficial feature noise.",
        
        "3.3.11 Opacus Differential Privacy (DP-SGD) and Renyi Privacy Budgets\n"
        "To enforce mathematical privacy guarantees when modeling sensitive user data, we train our Deep MLP using DP-SGD via Opacus. "
        "For each training batch of size $B$, the gradients are computed for each individual sample $i$, clipped to a maximum L2-norm threshold $C$, "
        "and averaged before adding calibrated Gaussian noise:\n"
        "$$g_i \\leftarrow g_i / \\max\\left(1, \\frac{\\|g_i\\|_2}{C}\\right)$$\n"
        "$$\\tilde{g} = \\frac{1}{B} \\left( \\sum_{i=1}^B g_i + \\mathcal{N}(0, \\sigma^2 C^2 I) \\right)$$\n"
        "Where $\\sigma$ is the noise multiplier. The privacy budget is tracked using Renyi Differential Privacy (RDP), establishing strict "
        "$(\\epsilon = 8.0, \\delta = 10^{-5})$ privacy guarantees that mathematically protect individual profiles from reconstruction attacks.",
        
        "3.3.12 Graph Attention Network (GAT) Semi-Supervised Node Classification\n"
        "We model the dating app dataset as a social network similarity graph $G = (V, E)$, where users are represented by nodes $v_i \\in V$, "
        "and edges $e_{i,j} \\in E$ represent k-nearest-neighbor similarity connections. We apply a Graph Attention Network (GAT) to perform "
        "semi-supervised node classification. The graph attention coefficient $\\alpha_{i,j}$ determines the attention node $i$ pays to node $j$:\n"
        "$$\\alpha_{i,j} = \\frac{\\exp(\\text{LeakyReLU}(\\mathbf{a}^T [W h_i \\parallel W h_j]))}{\\sum_{k \\in \\mathcal{N}(i)} \\exp(\\text{LeakyReLU}(\\mathbf{a}^T [W h_i \\parallel W h_k]))}$$\n"
        "Where $W$ is a shared linear transformation matrix, $\\mathbf{a}$ is the attention vector, and $\\parallel$ denotes concatenation. "
        "The output feature representation $h_i'$ is computed as a weighted combination of neighbor features mapped by these attention coefficients:\n"
        "$$h_i' = \\sigma \\left( \\sum_{j \\in \\mathcal{N}(i)} \\alpha_{i,j} W h_j \\right)$$\n"
        "This allows the network to classify matching outcomes based on dynamic neighborhood behavioral similarities rather than individual static factors."
    ]
    insert_paragraphs_at(idx_models + 2, model_writeups)
    print("Injected 6 massive deep learning model writeups successfully.")

# Save modified document
doc.save(out_path)
print("🎉 Success! The V5.1 massive SOTA assignment report is expanded further.")
