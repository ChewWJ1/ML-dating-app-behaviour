import docx

doc_path = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report.docx"
doc = docx.Document(doc_path)

for idx in range(178, min(184, len(doc.paragraphs))):
    print(f"[{idx}] {doc.paragraphs[idx].text}")
