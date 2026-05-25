import docx

doc_path = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report.docx"
doc = docx.Document(doc_path)

def print_range(start, end, label):
    print(f"\n=== {label} (indices {start} to {end}) ===")
    for idx in range(start, min(end, len(doc.paragraphs))):
        print(f"[{idx}] {doc.paragraphs[idx].text}")

print_range(70, 81, "Section 3.1")
print_range(153, 161, "Section 6.1")
print_range(182, 187, "Section 6.3")
