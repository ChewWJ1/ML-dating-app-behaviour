import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

root_dir = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour"
doc_path = os.path.join(root_dir, "reports", "WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx")

doc = docx.Document(doc_path)

print("Paragraph 201:")
print(doc.paragraphs[201].text)
print("\nParagraph 202:")
print(doc.paragraphs[202].text)
print("\nParagraph 203:")
print(doc.paragraphs[203].text)
print("\nParagraph 205:")
print(doc.paragraphs[205].text)
print("\nParagraph 206:")
print(doc.paragraphs[206].text)
print("\nParagraph 207:")
print(doc.paragraphs[207].text)
print("\nParagraph 397:")
print(doc.paragraphs[397].text)
print("\nParagraph 398:")
print(doc.paragraphs[398].text)

print("\nTable 10, Row 3, Cell 3:")
print(doc.tables[10].rows[3].cells[3].text)
