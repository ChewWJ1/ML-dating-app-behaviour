import docx
from docx import Document

doc = Document("reports/WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.1 SOTA.docx")
print(f"Total paragraphs in compiled SOTA doc: {len(doc.paragraphs)}")

found_refs = False
for idx, p in enumerate(doc.paragraphs):
    if "References" in p.text and len(p.text) < 20:
        print(f"Found References heading at index {idx}: '{p.text}'")
        found_refs = True
        # Print the next 50 paragraphs
        for next_idx in range(idx, min(idx + 50, len(doc.paragraphs))):
            text_to_print = doc.paragraphs[next_idx].text
            print(f"{next_idx}: '{text_to_print.encode('ascii', errors='replace').decode('ascii')}'")
        break

if not found_refs:
    print("Could not find 'References' heading in the compiled document.")
