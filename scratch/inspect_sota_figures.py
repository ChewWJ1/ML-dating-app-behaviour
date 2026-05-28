import docx
import os

doc_path = r"reports/WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.2 SOTA.docx"
if not os.path.exists(doc_path):
    print("Error: File not found!")
    exit(1)

doc = docx.Document(doc_path)
print(f"Total paragraphs in V5.2 SOTA: {len(doc.paragraphs)}")

for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "Figure " in text:
        print(f"P {idx}: {text}")
