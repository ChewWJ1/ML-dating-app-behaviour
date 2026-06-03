import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

DOCX_PATH = r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx'
doc = Document(DOCX_PATH)

print("=" * 70)
print("MAPPING FIGURES 38-42 AND TARGET SECTION HEADERS")
print("=" * 70)

# Find Figs 38-42 and all surrounding context
print("\n--- FIGURES 38-42 LOCATIONS ---")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if any(f"Figure {n}" in t for n in [38, 39, 40, 41, 42]):
        print(f"  Para {i}: {t}")

print("\n--- ALL SECTION HEADERS (look for 4.x, 5.x, 7.x, 8.x) ---")
section_keywords = [
    "4 Model", "4. Model", "Section 4", "4.1", "4.2", "4.3", "4.4",
    "5 Model", "5. Model", "Section 5", "5.1", "5.2", "5.3",
    "7 Implemented", "7. Implemented", "Section 7", "7.1", "7.2",
    "8 Conclusion", "8. Conclusion", "8.1", "8.2", "8.3",
    "Knowledge Distillation", "Model Compression", "Regularization",
    "Differential Privacy", "Self-Supervised", "SCARF", "Explainability",
    "Feature Attribution", "Advanced Model", "Robustness"
]
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t and any(kw.lower() in t.lower() for kw in section_keywords):
        # Only print if it looks like a heading (short-ish)
        if len(t) < 200:
            print(f"  Para {i} [style={p.style.name[:30]}]: {t}")

print("\n--- CONTEXT AROUND FIG 38-42 (paras around them) ---")
fig_paras = []
for i, p in enumerate(doc.paragraphs):
    if any(f"Figure {n}" in p.text for n in [38, 39, 40, 41, 42]):
        fig_paras.append(i)
        
for fp in fig_paras:
    start = max(0, fp - 2)
    end = min(len(doc.paragraphs), fp + 4)
    print(f"\n  === Around Figure at Para {fp} ===")
    for j in range(start, end):
        print(f"    Para {j}: {doc.paragraphs[j].text[:150]}")
