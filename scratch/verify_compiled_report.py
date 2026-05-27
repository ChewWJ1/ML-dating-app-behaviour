import docx
import os

doc_path = "reports/WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.1 SOTA.docx"
if not os.path.exists(doc_path):
    print("Document not found yet! Waiting for compilation to finish...")
    exit(1)

doc = docx.Document(doc_path)
print("=== Verification of Compiled Report ===")
print(f"File Path: {doc_path}")
print(f"Total Paragraphs: {len(doc.paragraphs)}")
print(f"Total Tables: {len(doc.tables)}")
print(f"Total Inline Shapes (Images): {len(doc.inline_shapes)}")

word_count = sum(len(p.text.split()) for p in doc.paragraphs)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            word_count += len(cell.text.split())
print(f"Total Estimated Word Count: {word_count}")

# Academic pages with figures and tables usually average 250-350 words per page.
# Each figure also adds about 1/3 to 1/2 of a page in Word layout.
# Let's estimate pages based on this logic.
est_text_pages = word_count / 300.0
est_img_pages = len(doc.inline_shapes) * 0.45
est_table_pages = len(doc.tables) * 0.4
total_est_pages = est_text_pages + est_img_pages + est_table_pages

print(f"Estimated Text Pages: {est_text_pages:.1f}")
print(f"Estimated Image Pages: {est_img_pages:.1f}")
print(f"Estimated Table Pages: {est_table_pages:.1f}")
print(f"Total Estimated Page Count in Word: {total_est_pages:.1f} pages")
