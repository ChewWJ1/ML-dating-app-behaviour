import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

DOCX_PATH = r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx'
doc = Document(DOCX_PATH)

print("=" * 70)
print("VERIFICATION OF ALL FIXES")
print("=" * 70)

# ── D1: Feature counts
print("\n[D1] Feature counts:")
print(f"  Para 143: {doc.paragraphs[143].text[:120]}")
print(f"  Para 150: {doc.paragraphs[150].text[:120]}")
print(f"  Para 373: {doc.paragraphs[373].text[:120]}")
print(f"  Para 26:  {doc.paragraphs[26].text[:120]}")

# Confirm no "116" or "67 features" remain
leftovers_116 = [(i, p.text) for i, p in enumerate(doc.paragraphs)
                  if "116 features" in p.text or "116 input features" in p.text]
leftovers_67f = [(i, p.text) for i, p in enumerate(doc.paragraphs) if "67 features" in p.text]
print(f"\n  Remaining '116 features' in paras: {len(leftovers_116)}")
for i, t in leftovers_116: print(f"    Para {i}: {t[:100]}")
print(f"  Remaining '67 features' in paras: {len(leftovers_67f)}")
for i, t in leftovers_67f: print(f"    Para {i}: {t[:100]}")

# ── D2: Engineered feature names
print("\n[D2] Engineered feature names:")
fake_names = ["popularity_density", "bio_message_interaction", "selective_emoji_swiper"]
leftovers_fake = [(i, p.text) for i, p in enumerate(doc.paragraphs)
                   if any(fn in p.text for fn in fake_names)]
real_names = [(i, p.text) for i, p in enumerate(doc.paragraphs)
               if any(rn in p.text for rn in ["engagement_score", "profile_completeness", "activity_intensity"])]
print(f"  Remaining FAKE name occurrences in paras: {len(leftovers_fake)}")
for i, t in leftovers_fake: print(f"    Para {i}: {t[:120]}")
print(f"  REAL name occurrences in paras: {len(real_names)}")
for i, t in real_names: print(f"    Para {i}: {t[:120]}")

# ── D3: DML p-value
print("\n[D3] DML p-value (exec summary):")
print(f"  Para 28: {doc.paragraphs[28].text[:200]}")

# ── D4: Gender parity
print("\n[D4] Gender parity:")
print(f"  Para 291: {doc.paragraphs[291].text[:200]}")

# ── D5: PCA
print("\n[D5] PCA components:")
print(f"  Para 153: {doc.paragraphs[153].text[:150]}")
t4r4c2 = doc.tables[4].rows[4].cells[2]
print(f"  Table 4 R4 C2: {t4r4c2.text[:150]}")

# ── B: Figure placement context
print("\n[B] Figure placement context:")
print(f"  Para 261 (5.2 intro): {doc.paragraphs[261].text[:180]}")
print(f"  Para 332 (7.1 intro): {doc.paragraphs[332].text[:180]}")
print(f"  Para 386 (before figs): {doc.paragraphs[386].text[-200:]}")
print(f"  Para 387 (new heading): {doc.paragraphs[387].text[:100]}")
print(f"  Para 220 (Fig 41 xref): {doc.paragraphs[220].text[-150:]}")

print("\n" + "=" * 70)
print("VERIFICATION COMPLETE")
print("=" * 70)
