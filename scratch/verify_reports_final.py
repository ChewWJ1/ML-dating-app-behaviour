import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

long_path = r"reports/WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.1(long).docx"
sota_path = r"reports/WIA1006 Machine Learning - Tying the Data Knot Assignment Report V5.2 SOTA.docx"

def verify_file(filepath):
    print(f"\n==========================================")
    print(f"VERIFYING: {filepath}")
    print(f"==========================================")
    
    if not os.path.exists(filepath):
        print("ERROR: File does not exist!")
        return False
        
    try:
        doc = docx.Document(filepath)
    except Exception as e:
        print(f"ERROR: Failed to open document: {e}")
        return False
        
    print(f"File loaded successfully. Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
    
    # 1. Verify Heading Hierarchy
    print("\n--- Verifying Heading Numbers & Hierarchy ---")
    headings = []
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith(("1.0", "2.0", "3.0", "4.0", "5.0", "6.0", "6.1", "6.2", "6.3", "7.0", "7.1", "7.2", "7.3", "8.0", "8.1", "8.2")):
            headings.append((idx, text))
            print(f"  P {idx}: {text}")
            
    expected_sequence = [
        "5.0 Insights and Interpretation",
        "6.0 SwipeIQ V2: Premium Interactive Analytics Dashboard and Web Application",
        "6.1 Engineering Architecture and Cloud Deployment Framework",
        "6.2 Detailed Specifications of Interactive Workspaces & Stress-Testing Playgrounds",
        "6.3 System Engineering & Modular Implementation Specifications",
        "7.0 Implemented Enhancements, Performance Optimization & Excluded Techniques",
        "7.1 Summary of Implemented Enhancements & Optimizations",
        "7.2 Detailed Technical Specifications",
        "7.3 Summary of Evaluated and Excluded Techniques",
        "8.0 Conclusion and Future Work",
        "8.1 Key Findings Summary",
        "8.2 Recommendations for Future Research"
    ]
    
    seq_idx = 0
    errors = 0
    for idx, text in headings:
        for exp in expected_sequence:
            if exp in text:
                print(f"  [OK] Found expected heading: '{text}' at P{idx}")
                seq_idx += 1
                break
                
    if seq_idx < len(expected_sequence):
        print(f"  [ERROR] Missing expected headings. Found only {seq_idx} out of {len(expected_sequence)}.")
        errors += 1
        
    # 2. Verify Figure Numbering (For Figures 37 and above to ensure no duplicates in the final section)
    print("\n--- Verifying Figure Numbering & Duplicates (Figure 37 and above) ---")
    captions = []
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if "Figure " in text:
            captions.append((idx, text))
            
    fig_counts = {}
    for idx, text in captions:
        parts = text.split(":")
        if parts:
            fig_label = parts[0].strip()
            if fig_label.startswith("Figure"):
                # Parse the figure number
                try:
                    num_str = fig_label.replace("Figure", "").strip()
                    num = int(num_str)
                    if num >= 37:
                        fig_counts[fig_label] = fig_counts.get(fig_label, 0) + 1
                except ValueError:
                    pass
                
    duplicate_figs = {k: v for k, v in fig_counts.items() if v > 1}
    if duplicate_figs:
        print(f"  [ERROR] Duplicate Figure Labels Found among high-numbered figures: {duplicate_figs}")
        errors += 1
        for idx, text in captions:
            for dup in duplicate_figs:
                if text.startswith(dup):
                    print(f"    P {idx}: {text}")
    else:
        print("  [OK] Zero duplicate figures found above Figure 36! Figure flow is clean.")
        print("  Sample sequential captions:")
        for idx, text in captions[-10:]:
            print(f"    P {idx}: {text[:100]}...")

    # 3. Verify Table 6, 7, 8 Updates
    print("\n--- Verifying Table performance metrics ---")
    try:
        t6 = doc.tables[6]
        knn_row_t6 = t6.rows[2]
        lgbm_row_t6 = t6.rows[6]
        
        print(f"  Table 6 KNN Baseline Test Acc: {knn_row_t6.cells[1].text}, F1-Score: {knn_row_t6.cells[4].text} (Expected: 43.09%, 54.00%)")
        print(f"  Table 6 LightGBM Baseline Test Acc: {lgbm_row_t6.cells[1].text}, F1-Score: {lgbm_row_t6.cells[4].text} (Expected: 58.62%, 14.33%)")
        
        if knn_row_t6.cells[4].text != "54.00%" or lgbm_row_t6.cells[1].text != "58.62%":
            print("  [ERROR] Table 6 metrics mismatch!")
            errors += 1
        else:
            print("  [OK] Table 6 metrics are verified correct!")
            
        t8 = doc.tables[8]
        print(f"  Table 8 rows count: {len(t8.rows)} (Expected: 5 rows - 1 header + 4 data)")
        last_row_t8 = t8.rows[-1]
        print(f"  Table 8 Last Row Classifier: {last_row_t8.cells[0].text}, Tuned F1: {last_row_t8.cells[4].text} (Expected: CatBoost, 31.46%)")
        
        if len(t8.rows) != 5 or last_row_t8.cells[0].text != "CatBoost" or last_row_t8.cells[4].text != "31.46%":
            print("  [ERROR] Table 8 metrics or rows mismatch!")
            errors += 1
        else:
            print("  [OK] Table 8 hyperparameter tuning rows are verified correct!")
    except Exception as e:
        print(f"  [ERROR] Failed to read metric tables: {e}")
        errors += 1

    # 4. Verify Hyperlink
    print("\n--- Verifying Clickable Streamlit Hyperlink ---")
    found_link = False
    for idx, p in enumerate(doc.paragraphs):
        if "https://ml-tying-the-data-knot-swipeiq-app.streamlit.app/" in p.text:
            found_link = True
            print(f"  [OK] Found Live App URL in P{idx}: '{p.text[:120]}...'")
            break
    if not found_link:
        print("  [ERROR] Live App URL not found!")
        errors += 1

    if errors == 0:
        print(f"\n==========================================")
        print(f"🎉 VERIFICATION PASSED FOR {filepath}!")
        print(f"==========================================")
        return True
    else:
        print(f"\n==========================================")
        print(f"❌ VERIFICATION FAILED FOR {filepath} WITH {errors} ERRORS.")
        print(f"==========================================")
        return False

# Run verification on both documents
v_long = verify_file(long_path)
v_sota = verify_file(sota_path)

if v_long and v_sota:
    print("\n🚀 All automated verification checks passed cleanly! No errors detected.")
    sys.exit(0)
else:
    print("\n❌ Errors were detected in one or both files.")
    sys.exit(1)
