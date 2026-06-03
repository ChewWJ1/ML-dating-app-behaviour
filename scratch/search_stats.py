import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx')

print("=== ALL OCCURRENCES OF 'friedman' ===")
for i, p in enumerate(doc.paragraphs):
    if 'friedman' in p.text.lower():
        print(f"Para {i}: {p.text}")
        print("-" * 50)
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if 'friedman' in cell.text.lower():
                print(f"Table {t_idx}, Row {r_idx}, Col {c_idx}: {cell.text}")
                print("-" * 50)

print("=== ALL OCCURRENCES OF 't-test' OR 'paired' ===")
for i, p in enumerate(doc.paragraphs):
    if 't-test' in p.text.lower() or 'paired' in p.text.lower():
        print(f"Para {i}: {p.text}")
        print("-" * 50)
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if 't-test' in cell.text.lower() or 'paired' in cell.text.lower():
                print(f"Table {t_idx}, Row {r_idx}, Col {c_idx}: {cell.text}")
                print("-" * 50)
