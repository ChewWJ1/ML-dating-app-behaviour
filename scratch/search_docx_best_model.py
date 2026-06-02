import docx
import re
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx"
doc = docx.Document(doc_path)

print("Searching for 'best model' or similar in paragraphs...")

keywords = [r"best model", r"best-performing model", r"champion model", r"selected model", r"best performing model"]

for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    for kw in keywords:
        if re.search(kw, text, re.IGNORECASE):
            print(f"\nP{idx}: {text[:250]}...")
            break
            
print("\nSearching for 'best model' or similar in tables...")
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text.strip().replace("\n", " ")
            for kw in keywords:
                if re.search(kw, text, re.IGNORECASE):
                    print(f"Table {t_idx} Row {r_idx} Col {c_idx}: {text[:150]}")
                    break
