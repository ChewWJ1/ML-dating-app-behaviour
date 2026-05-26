import docx

doc_path = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report.docx"
doc = docx.Document(doc_path)

with open("scratch/docx_full_text.txt", "w", encoding="utf-8") as f:
    for i, para in enumerate(doc.paragraphs):
        f.write(f"[{i}] {para.text}\n")

print(f"Full text written. Total paragraphs: {len(doc.paragraphs)}")
