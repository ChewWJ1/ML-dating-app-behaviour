"""
Comprehensive fix script for WIA1006 ML Report V8 (final).docx
Fixes all remaining inconsistencies found during the deep audit:
1. StandardScaler -> RobustScaler in Para 105 (preprocessing list)
2. "Random Forest" -> "Dynamic Champion Model (LightGBM)" in Figure 18 description (Para 164)
3. "XGBoost model" -> "Dynamic Champion Model" in ipywidgets section (Para 389)
4. "113 columns" -> "116 features" in conclusion (Para 397)
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
import copy
import re

DOCX_PATH = r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx'

doc = Document(DOCX_PATH)

changes_made = []

def replace_in_para(para, old_text, new_text, para_idx):
    """Replace text in a paragraph across all runs, preserving formatting of first run."""
    full_text = ''.join(r.text for r in para.runs)
    if old_text not in full_text:
        return False
    
    new_full_text = full_text.replace(old_text, new_text, 1)
    
    # Clear all runs and set text in first run
    if para.runs:
        para.runs[0].text = new_full_text
        for run in para.runs[1:]:
            run.text = ''
    
    changes_made.append(f"Para {para_idx}: '{old_text[:60]}' -> '{new_text[:60]}'")
    return True

# ============================================================
# FIX 1: Para 105 - StandardScaler -> RobustScaler
# ============================================================
p105 = doc.paragraphs[105]
old = 'Normalization: Applied a StandardScaler to all 12 numerical features (centering to mean=0 and scaling to unit variance). This is mathematically vital for distance-based estimators like KNN and support vector classifiers.'
new = 'Normalization: Applied a RobustScaler to all 12 numerical features (centering using the median and scaling using the Interquartile Range). This is mathematically vital for distance-based estimators like KNN and support vector classifiers, as it is resistant to outlier distortions from power users.'

if old in p105.text:
    replace_in_para(p105, old, new, 105)
    print(f'[FIX 1] Para 105: StandardScaler -> RobustScaler DONE')
else:
    print(f'[FIX 1] Para 105: Not found, checking text: {p105.text[:200]}')

# ============================================================
# FIX 2: Para 164 - Figure 18 description "selected best model (Random Forest)"
# -> "Dynamic Champion Model (LightGBM)"
# ============================================================
p164 = doc.paragraphs[164]
old_164 = 'The selected best model (Random Forest) is calibrated via Isotonic Regression, explained via SHAP TreeExplainer, and deployed to generate counterfactual recourse recommendations (DiCE) and causal treatment uplifts (T-learner) for targeted recommendations.'
new_164 = 'The Dynamic Champion Model (LightGBM, dynamically selected by highest ROC-AUC) is calibrated via Isotonic Regression, explained via SHAP TreeExplainer, and deployed to generate counterfactual recourse recommendations (DiCE) and causal treatment uplifts (T-learner) for targeted recommendations.'

if old_164 in p164.text:
    replace_in_para(p164, old_164, new_164, 164)
    print(f'[FIX 2] Para 164: "Random Forest" in Figure 18 description -> "Dynamic Champion Model (LightGBM)" DONE')
else:
    print(f'[FIX 2] Para 164: Old text not found, checking: {p164.text[:300]}')

# ============================================================
# FIX 3: Para 389 - ipywidgets "XGBoost model" -> "Dynamic Champion Model"
# ============================================================
p389 = doc.paragraphs[389]
old_389 = 'evaluates the trained XGBoost model in real-time'
new_389 = 'evaluates the trained Dynamic Champion Model in real-time'

if old_389 in p389.text:
    replace_in_para(p389, old_389, new_389, 389)
    print(f'[FIX 3] Para 389: "XGBoost model" in simulator -> "Dynamic Champion Model" DONE')
else:
    print(f'[FIX 3] Para 389: Not found, checking: {p389.text[:300]}')

# ============================================================
# FIX 4: Para 397 - "113 columns" -> "116 features" (correct V8 count)
# ============================================================
p397 = doc.paragraphs[397]
old_397 = 'expanding features from 25 to 113 columns through ordinal, one-hot, and multi-hot encodings'
new_397 = 'expanding features from 25 to 116 features through ordinal, one-hot, and multi-hot encodings (including 3 engineered interaction features)'

if old_397 in p397.text:
    replace_in_para(p397, old_397, new_397, 397)
    print(f'[FIX 4] Para 397: "113 columns" -> "116 features" DONE')
else:
    print(f'[FIX 4] Para 397: Not found, checking: {p397.text[:300]}')

# ============================================================
# SAVE
# ============================================================
doc.save(DOCX_PATH)

print()
print(f'=== FIXES APPLIED: {len(changes_made)} ===')
for c in changes_made:
    print(f'  ✓ {c}')
print()
print('Document saved successfully.')
