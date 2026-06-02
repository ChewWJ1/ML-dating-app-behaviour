import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx')

print("=== ALL PARAGRAPHS 394 TO 406 ===")
for i in range(394, min(407, len(doc.paragraphs))):
    print(f"Para {i}: {doc.paragraphs[i].text}")
    print("-" * 50)
