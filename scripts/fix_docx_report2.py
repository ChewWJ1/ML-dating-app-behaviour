import docx
import os
import re

doc_path = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx"
out_path = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx"

print(f"Loading document: {doc_path}")
doc = docx.Document(doc_path)

replacements = [
    # Case-insensitive replacements for the missed items
    (r"Isotonic regression successfully calibrated the XGBoost \(Tuned\) champion", 
     r"Isotonic regression successfully calibrated the Dynamic Champion Model"),
    (r"the XGBoost \(Tuned\) algorithm is essentially forced", r"the Champion algorithm is essentially forced")
]

def replace_text_in_runs(paragraph, replacements):
    full_text = paragraph.text
    original_text = full_text
    
    for old, new in replacements:
        if re.search(old, full_text, re.IGNORECASE):
            full_text = re.sub(old, new, full_text, flags=re.IGNORECASE)
            
    if full_text != original_text:
        if len(paragraph.runs) > 0:
            style = paragraph.runs[0].style
            bold = paragraph.runs[0].bold
            italic = paragraph.runs[0].italic
            underline = paragraph.runs[0].underline
            
            for run in paragraph.runs:
                run.text = ""
                
            paragraph.runs[0].text = full_text
            paragraph.runs[0].bold = bold
            paragraph.runs[0].italic = italic
            paragraph.runs[0].underline = underline

# Patch Paragraphs
for i, para in enumerate(doc.paragraphs):
    if i == 255 and "LightGBM:" in para.text:
        # Directly replace the specific paragraph 255
        para.text = para.text.replace("LightGBM:", "Random Forest:")
    
    replace_text_in_runs(para, replacements)

# Patch Tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                replace_text_in_runs(para, replacements)

doc.save(out_path)
print(f"Successfully patched remaining cases and saved to {out_path}")
