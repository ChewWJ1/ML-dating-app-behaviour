import docx
import os
import re

doc_path = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx"
out_path = r"c:\Users\HP\Documents\GitHub\ML-dating-app-behaviour\reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx"

print(f"Loading document: {doc_path}")
doc = docx.Document(doc_path)

replacements = [
    # Exec Summary & General
    (r"we formally select the Isotonically Calibrated XGBoost \(Tuned\) as our final best-performing model", 
     r"we formally utilize a Dynamic Champion Model (dynamically inheriting the best weights based on ROC-AUC, e.g., LightGBM) as our final best-performing model"),
    (r"the XGBoost \(Tuned\) as our final best model", r"the Dynamic Champion Model as our final best model"),
    (r"highest individual test accuracy of 39\.71%", r"highest individual test accuracy of 60.48%"),
    
    # Section 3
    (r"Sixteen baseline and advanced classifiers—spanning", r"14 custom architectures and 2 AutoML baselines (total 16 models)—spanning"),
    (r"Sixteen baseline and advanced classifiers", r"14 custom architectures and 2 AutoML baselines (total 16 models)"),
    
    # Section 4.1.2
    (r"Random Forest is designated as the final selected best model for the following reasons:", 
     r"The Dynamic Champion Model is designated as the final selected best model for the following reasons:"),
    (r"Random Forest achieves a test accuracy", r"The Champion Model achieves a test accuracy"),
    (r"Random Forest provides native compatibility with SHAP TreeExplainer", 
     r"The tree-based champion model provides native compatibility with SHAP TreeExplainer"),
    (r"The calibrated Random Forest directly powers", r"The calibrated Champion Model directly powers"),
    (r"a requirement satisfied by Random Forest but not by", r"a requirement satisfied by standard tree ensembles but not by"),
    (r"By aggregating 500 independent decision trees trained on bootstrapped subsets, Random Forest stabilizes predictions", 
     r"By aggregating independent decision trees trained on bootstrapped subsets, the Champion Model stabilizes predictions"),
    (r"matches that of the Random Forest", r"matches that of the Champion Model"),
    
    # Section 4.3 (Confusion Matrix & Best Model)
    (r"Figure 28: Detailed Confusion Matrix of the Selected Best Model \(XGBoost \(Tuned\)\)", 
     r"Figure 28: Detailed Confusion Matrix of the Selected Best Model"),
    
    # Section 5.2
    (r"our selected best model \(XGBoost \(Tuned\)\)", r"our selected best model"),
    (r"The XGBoost \(Tuned\) algorithm is essentially forced", r"The Champion algorithm is essentially forced"),
    
    # Section 8.1
    (r"formally selected the XGBoost \(Tuned\) as our final best model", r"formally utilized a Dynamic Champion Model as our final best model"),
    (r"the XGBoost \(Tuned\) was selected because its tree-based ensemble architecture enabled", 
     r"the Dynamic Champion Model was selected because its tree-based ensemble architecture enabled"),
]

def replace_text_in_runs(paragraph, replacements):
    # Because text might be split across multiple runs in python-docx, 
    # we'll do a simple trick: if the full paragraph text contains the pattern, 
    # we replace it in the first run and clear the other runs. This loses some internal 
    # formatting (like bolding a single word inside the sentence) but preserves the 
    # paragraph's overall style. 
    # A safer way if it spans runs is to just concatenate and put in run[0].
    
    full_text = paragraph.text
    original_text = full_text
    
    for old, new in replacements:
        if re.search(old, full_text):
            full_text = re.sub(old, new, full_text)
            
    if full_text != original_text:
        # Clear all runs and put new text in the first run
        if len(paragraph.runs) > 0:
            # Preserve the formatting of the first run
            style = paragraph.runs[0].style
            bold = paragraph.runs[0].bold
            italic = paragraph.runs[0].italic
            underline = paragraph.runs[0].underline
            font_name = paragraph.runs[0].font.name
            font_size = paragraph.runs[0].font.size
            
            for run in paragraph.runs:
                run.text = ""
                
            paragraph.runs[0].text = full_text
            # Restore formatting
            paragraph.runs[0].bold = bold
            paragraph.runs[0].italic = italic
            paragraph.runs[0].underline = underline
            
print("Patching paragraphs...")
# Patch Paragraphs
lgbm_count = 0
for para in doc.paragraphs:
    # Handle the duplicate LightGBM hyperparameter grid issue
    if "LightGBM:" in para.text:
        lgbm_count += 1
        if lgbm_count == 1 and "min_samples_split" in para.text:
            # This is the Random Forest grid mislabeled as LightGBM
            para.text = para.text.replace("LightGBM:", "Random Forest:")
    
    replace_text_in_runs(para, replacements)

print("Patching tables...")
# Patch Tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                replace_text_in_runs(para, replacements)

doc.save(out_path)
print(f"Successfully patched and saved to {out_path}")
