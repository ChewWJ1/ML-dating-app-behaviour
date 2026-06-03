"""
Definitive fix: For each figure block, collect all elements in that block,
sort them into the correct order, and reinsert atomically.
No reversed() — just insert in forward order using addprevious on a fresh anchor.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from lxml import etree

DOCX_PATH = r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx'
doc = Document(DOCX_PATH)

def has_image(p):
    xml = etree.tostring(p._element, encoding='unicode')
    return '<w:drawing' in xml or '<v:shape' in xml

def show(paras, a, b):
    for i in range(a, min(b, len(paras))):
        p = paras[i]
        img = '[IMG]' if has_image(p) else '     '
        print(f"  {i:4d} {img} {p.text[:108]}")

def move_block_before(doc, ordered_indices, anchor_text):
    """
    Move the paragraphs at ordered_indices to appear (in that order) 
    immediately before the paragraph containing anchor_text.
    """
    paras = doc.paragraphs
    # Find anchor
    anchor_elem = None
    for p in paras:
        if anchor_text in p.text:
            anchor_elem = p._element
            break
    if anchor_elem is None:
        print(f"  ERROR: anchor not found: '{anchor_text[:60]}'")
        return
    # Collect elements
    elems = [paras[i]._element for i in ordered_indices]
    # Remove from current location
    for e in elems:
        e.getparent().remove(e)
    # Insert in correct forward order: iterate forward, each one goes before anchor
    # To maintain order when inserting before anchor: first insert last element, then second-to-last, etc.
    # Actually: insert elem[0] before anchor, then insert elem[1] AFTER elem[0] = before anchor still works
    # Simplest approach: insert all before anchor in reverse order -> they pile up correctly
    for e in reversed(elems):
        anchor_elem.addprevious(e)
    print(f"  Moved {ordered_indices} -> before '{anchor_text[:50]}'")

# ══════════════════════════════════════════════════════════════════════════════
# Current state check
paras = doc.paragraphs
print("CURRENT STATE:")

print("\n== Fig 41 block ==")
show(paras, 233, 241)

print("\n== Fig 42 block ==")
show(paras, 262, 269)

print("\n== Figs 39+40 block ==")
show(paras, 288, 300)

print("\n== Fig 38 block ==")
show(paras, 385, 393)

print("\n══════════ APPLYING FIXES ══════════\n")

# ── FIG 41: Currently body(235) -> image(236) -> caption(237) -> intro(238)
# Target: intro(238) -> caption(237) -> image(236) -> body(235)
# Current indices: 235=body, 236=image, 237=caption, 238=intro
# Order we want when reading down the doc: intro, caption, image, body
move_block_before(doc, [238, 237, 236, 235], "4.3 Hyperparameter Tuning")
doc.save(DOCX_PATH)
doc = Document(DOCX_PATH); paras = doc.paragraphs

print("\nFig 41 after fix:")
show(paras, 233, 241)

# ── FIG 42: Currently body(264) -> caption(265) -> intro(266)
# Target: intro -> caption -> body (no image found)
move_block_before(doc, [266, 265, 264], "5.0 Insights and Interpretation")
doc.save(DOCX_PATH)
doc = Document(DOCX_PATH); paras = doc.paragraphs

print("\nFig 42 after fix:")
show(paras, 262, 270)

# ── FIGs 39+40: Current state from final_reorder output:
#   290=disclosure40, 291=body40, 292=image(39 or 40), 293=caption40
#   294=body39, 295=image39, 296=caption39, 297=intro
# Target order: intro, cap39, img39, body39, cap40, img40, body40, disclosure40
# 
# Need to re-identify after previous moves
doc = Document(DOCX_PATH); paras = doc.paragraphs
f39 = {}; f40 = {}; intro_39_40 = None
for i, p in enumerate(paras):
    t = p.text
    if "Complementing the SHAP attributions above" in t:       intro_39_40 = i
    elif "Figure 39: Attentive TabNet" in t:                   f39['caption'] = i
    elif "attentive feature selection heatmap in Figure 39" in t: f39['body'] = i
    elif "Figure 40: SCARF Self-Supervised" in t:              f40['caption'] = i
    elif "Figure 40 projects the SCARF contrastive" in t:      f40['body'] = i
    elif "Methodological Disclosure: During our SCARF" in t:   f40['disclosure'] = i

# Find images by proximity to known captions
all_blank_imgs = [i for i, p in enumerate(paras) if has_image(p) and p.text.strip() == '']
# Pick the two images closest to intro_39_40
if intro_39_40:
    nearby = sorted(all_blank_imgs, key=lambda x: abs(x - intro_39_40))
    nearby_close = [x for x in nearby if abs(x - intro_39_40) <= 10][:2]
    if len(nearby_close) >= 2:
        # The one closer to cap39 goes to f39, the other to f40
        if 'caption' in f39 and 'caption' in f40:
            d39 = [abs(x - f39['caption']) for x in nearby_close]
            d40 = [abs(x - f40['caption']) for x in nearby_close]
            if min(d39) < min(d40):
                f39['image'] = nearby_close[d39.index(min(d39))]
                remaining = [x for x in nearby_close if x != f39['image']]
                if remaining: f40['image'] = remaining[0]
            else:
                f40['image'] = nearby_close[d40.index(min(d40))]
                remaining = [x for x in nearby_close if x != f40['image']]
                if remaining: f39['image'] = remaining[0]
    elif len(nearby_close) == 1:
        f39['image'] = nearby_close[0]

print(f"\nFigs 39+40: intro={intro_39_40}, f39={f39}, f40={f40}")
if intro_39_40 is not None:
    show(paras, intro_39_40 - 3, max([v for v in list(f39.values())+list(f40.values())+[intro_39_40] if isinstance(v,int)]) + 3)

# Build ordered list
order_3940 = []
if intro_39_40 is not None:              order_3940.append(intro_39_40)
if 'caption' in f39:                     order_3940.append(f39['caption'])
if 'image' in f39:                       order_3940.append(f39['image'])
if 'body' in f39:                        order_3940.append(f39['body'])
if 'caption' in f40:                     order_3940.append(f40['caption'])
if 'image' in f40:                       order_3940.append(f40['image'])
if 'body' in f40:                        order_3940.append(f40['body'])
if 'disclosure' in f40:                  order_3940.append(f40['disclosure'])

print("Order:", order_3940)
move_block_before(doc, order_3940, "5.3 Demographic Parity")
doc.save(DOCX_PATH)
doc = Document(DOCX_PATH); paras = doc.paragraphs

print("\nFigs 39+40 after fix:")
for i, p in enumerate(paras):
    if any(kw in p.text for kw in ["Figure 39", "Figure 40", "AttentiveTab", "SCARF", "Complementing", "Methodological Disclosure: During", "5.3 Demographic"]):
        img = '[IMG]' if has_image(p) else '     '
        print(f"  {i:4d} {img} {p.text[:105]}")

# ── FIG 38
doc = Document(DOCX_PATH); paras = doc.paragraphs
f38 = {}; intro_38 = None
for i, p in enumerate(paras):
    t = p.text
    if "Figure 38 below presents the training loss curves comparing" in t: intro_38 = i
    elif "Figure 38: Knowledge Distillation Student" in t: f38['caption'] = i
    elif "Figure 38 compares the training loss profiles of the teacher ensemble" in t: f38['body'] = i
all_blank_imgs = [i for i, p in enumerate(paras) if has_image(p) and p.text.strip() == '']
if 'caption' in f38:
    close = [i for i in all_blank_imgs if abs(i - f38['caption']) <= 5]
    if close: f38['image'] = min(close, key=lambda x: abs(x - f38['caption']))

print(f"\nFig38: intro={intro_38}, f38={f38}")
if intro_38 is not None:
    show(paras, intro_38 - 2, max([v for v in list(f38.values())+[intro_38] if isinstance(v,int)])+3)

order_38 = []
if intro_38 is not None:   order_38.append(intro_38)
if 'caption' in f38:       order_38.append(f38['caption'])
if 'image' in f38:         order_38.append(f38['image'])
if 'body' in f38:          order_38.append(f38['body'])

print("Order:", order_38)
move_block_before(doc, order_38, "7.3 Summary of Evaluated and Excluded")
doc.save(DOCX_PATH)
doc = Document(DOCX_PATH); paras = doc.paragraphs

print("\nFig 38 after fix:")
for i, p in enumerate(paras):
    if any(kw in p.text for kw in ["Figure 38", "Knowledge Distill", "StudentNet", "7.3 Summary"]):
        img = '[IMG]' if has_image(p) else '     '
        print(f"  {i:4d} {img} {p.text[:105]}")

print("\n══════════ COMPLETE ══════════")
print("Document saved.")
