"""
Comprehensive audit script to verify that no remaining traces of the reported errors
exist anywhere in the document (paragraphs or tables).
"""

import sys
import re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

DOCX_PATH = r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx'
doc = Document(DOCX_PATH)

print("=== RUNNING SYSTEM-WIDE AUDIT FOR SPECIFIC REPORT PROBLEMS ===")

issues_found = []

# 1. Check for any remaining occurrences of "16 models" or "14 models" (case-insensitive)
# We exclude correct references (like citation #14 and #16, 14 surgical fixes, 14-page dashboards)
model_count_pattern = re.compile(r'\b(14|16)\s+(baseline|advanced|diverse|custom|classification|stacking)?\s*models?\b', re.IGNORECASE)

# 2. Check for any remaining occurrences of "5-fold" or "5 fold" CV/validation/t-test references
fold_pattern = re.compile(r'\b5-?folds?\b|\b5\s+validation\s+folds\b|\b5\s+independent\s+validation\b', re.IGNORECASE)

# 3. Check for any remaining occurrences of "logistic regression student" or variants
student_pattern = re.compile(r'logistic\s+regression\s+student|student\s+logistic\s+regression|logistic\s+student|student\s+logistic', re.IGNORECASE)

# Audit Paragraphs
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if not txt:
        continue
    
    # Check model count
    for m in model_count_pattern.finditer(txt):
        match_str = m.group()
        # Exclude correct phrases like "14 surgical fixes"
        if "fix" not in match_str.lower() and "dashboard" not in match_str.lower() and "page" not in match_str.lower():
            issues_found.append(f"Para {i} (model count): Found '{match_str}' in:\n  \"{txt[:120]}...\"")
            
    # Check folds
    for m in fold_pattern.finditer(txt):
        match_str = m.group()
        issues_found.append(f"Para {i} (fold count): Found '{match_str}' in:\n  \"{txt[:120]}...\"")
        
    # Check student
    for m in student_pattern.finditer(txt):
        match_str = m.group()
        issues_found.append(f"Para {i} (student model): Found '{match_str}' in:\n  \"{txt[:120]}...\"")

# Audit Tables
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        seen_cells = set()
        for c_idx, cell in enumerate(row.cells):
            if cell in seen_cells:
                continue
            seen_cells.add(cell)
            txt = cell.text.strip()
            if not txt:
                continue
            
            # Check model count
            for m in model_count_pattern.finditer(txt):
                match_str = m.group()
                if "fix" not in match_str.lower() and "dashboard" not in match_str.lower() and "page" not in match_str.lower():
                    issues_found.append(f"Table {t_idx}, R{r_idx}, C{c_idx} (model count): Found '{match_str}' in:\n  \"{txt[:120]}...\"")
                    
            # Check folds
            for m in fold_pattern.finditer(txt):
                match_str = m.group()
                issues_found.append(f"Table {t_idx}, R{r_idx}, C{c_idx} (fold count): Found '{match_str}' in:\n  \"{txt[:120]}...\"")
                
            # Check student
            for m in student_pattern.finditer(txt):
                match_str = m.group()
                issues_found.append(f"Table {t_idx}, R{r_idx}, C{c_idx} (student model): Found '{match_str}' in:\n  \"{txt[:120]}...\"")

print()
if issues_found:
    print(f"⚠️ AUDIT WARNING: Found {len(issues_found)} potentially unresolved issues in the report:")
    for issue in issues_found:
        print(f"  - {issue}")
        print("-" * 50)
else:
    print("✨ ALL CHECKS PASSED! No incorrect model counts, fold counts, or student model descriptions were found anywhere in the document.")
