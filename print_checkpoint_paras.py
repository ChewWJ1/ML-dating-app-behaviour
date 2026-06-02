import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx')

print("=== ALL PARAGRAPHS 373 TO 385 ===")
for i in range(373, min(386, len(doc.paragraphs))):
    print(f"Para {i}: {doc.paragraphs[i].text}")
    print("-" * 50)
