"""
Physically relocate Figures 38-42 from Section 8 to their correct sections.
Uses lxml XML manipulation to move paragraph elements within the document body.

Move plan (all indices are ORIGINAL, pre-move indices):
  Fig 41 (Mixup BCE):      paras 398-400  -> before para 233 (4.3 Hyperparameter Tuning)
  Fig 42 (Opacus DP-SGD): paras 401-402  -> before para 257 (5.0 Insights)
  Figs 39+40 (TabNet+SCARF): paras 391-397 -> before para 279 (5.3 Demographic Parity)
  Fig 38 (Knowledge Dist): paras 388-390  -> before para 367 (7.3 Excluded Techniques)

Also:
  - Remove the fake "8.3 Supplementary" heading (para 387) added in previous fix
  - Remove cross-reference sentences injected previously in paras 261, 332, 220 (too)
  - Add smooth transitional text at each insertion point
  - Fix the body text surrounding para 386 (active learning) by removing the injected sentence
  - Clean up para 403 (future recommendations text) since it immediately followed fig 42
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import copy

DOCX_PATH = r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx'
doc = Document(DOCX_PATH)

body = doc.element.body
all_paras = doc.paragraphs  # snapshot — do NOT use after mutations

# ── Helper: get body child index of a paragraph element ──────────────────────
def body_index(para):
    for i, child in enumerate(body):
        if child is para._element:
            return i
    return -1

# ── Helper: collect consecutive paragraph elements by original para indices ──
def collect_elems(para_indices):
    return [doc.paragraphs[i]._element for i in para_indices]

# ── Helper: insert list of elements before a reference element ────────────────
def insert_before(ref_elem, elems):
    for e in elems:
        ref_elem.addprevious(e)

# ── Helper: insert a new plain paragraph before a reference element ───────────
from docx.oxml import OxmlElement
def new_para_elem(text, bold=False, italic=False):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    if bold or italic:
        rPr = OxmlElement('w:rPr')
        if bold:
            b = OxmlElement('w:b'); rPr.append(b)
        if italic:
            i = OxmlElement('w:i'); rPr.append(i)
        r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p

def blank_para_elem():
    p = OxmlElement('w:p')
    return p

print("Reading original paragraph elements...")

# ─────────────────────────────────────────────────────────────────────────────
# STEP 1: Collect all elements we'll be moving (by original index, before any moves)
# Figure 38: caption(388) + body(389) + image(390)
# Figure 39: caption(391) + body(392) + image(393)
# Figure 40: caption(394) + body(395) + disclosure(396) + image(397)  
# Figure 41: caption(398) + body(399) + image(400)
# Figure 42: caption(401) + body(402)
# Fake 8.3 heading: (387)
# Also blank para before fig38 that we need to clean up

fig38_paras = [388, 389, 390]    # Knowledge Distillation
fig39_paras = [391, 392, 393]    # AttentiveTabNet
fig40_paras = [394, 395, 396, 397]  # SCARF (includes disclosure + image)
fig41_paras = [398, 399, 400]    # Mixup BCE
fig42_paras = [401, 402]         # Opacus DP-SGD

heading_8_3 = [387]              # The "8.3 Supplementary..." heading to remove

# Collect elements NOW before any mutation
fig38_elems  = collect_elems(fig38_paras)
fig39_elems  = collect_elems(fig39_paras)
fig40_elems  = collect_elems(fig40_paras)
fig41_elems  = collect_elems(fig41_paras)
fig42_elems  = collect_elems(fig42_paras)
heading_elem = doc.paragraphs[387]._element

# Anchor reference elements (where to insert before)
anchor_4_3   = doc.paragraphs[233]._element   # "4.3 Hyperparameter Tuning"
anchor_5_0   = doc.paragraphs[257]._element   # "5.0 Insights and Interpretation"
anchor_5_3   = doc.paragraphs[279]._element   # "5.3 Demographic Parity"
anchor_7_3   = doc.paragraphs[367]._element   # "7.3 Summary of Excluded Techniques"

# ─────────────────────────────────────────────────────────────────────────────
# STEP 2: Clean up cross-reference injections from previous fix script
# Para 261: "In addition to SHAP TreeExplainer attributions, the AttentiveTabNet..."
p261 = doc.paragraphs[261]
old_261 = ''.join(r.text for r in p261.runs)
if "AttentiveTabNet feature selection heatmap (Figure 39)" in old_261:
    # This whole paragraph was injected - blank it out so it becomes a simple spacer
    for r in p261.runs:
        r.text = ''
    print("  Cleared injected para 261 (5.2 cross-ref)")

# Para 332: "...Figure 38 (Section 8.3) visualizes..."
p332 = doc.paragraphs[332]
old_332 = ''.join(r.text for r in p332.runs)
injected_332 = " Figure 38 (Section 8.3) visualizes the knowledge distillation training loss curves that directly support the compression enhancement:"
if injected_332 in old_332:
    new_332 = old_332.replace(injected_332, ":")
    if p332.runs:
        p332.runs[0].text = new_332
        for r in p332.runs[1:]: r.text = ''
    print("  Cleaned para 332 (7.1 cross-ref)")

# Para 220: BCE/Mixup cross-ref
p220 = doc.paragraphs[220]
old_220 = ''.join(r.text for r in p220.runs)
injected_220 = " Figure 41 (Section 8.3) plots the BCE vs. Label-Smoothed Mixup loss curves empirically validating this regularization."
if injected_220 in old_220:
    new_220 = old_220.replace(injected_220, "")
    if p220.runs:
        p220.runs[0].text = new_220
        for r in p220.runs[1:]: r.text = ''
    print("  Cleaned para 220 (Fig 41 cross-ref)")

# Para 386: remove the injected transitional sentence
p386 = doc.paragraphs[386]
old_386 = ''.join(r.text for r in p386.runs)
injected_386 = (" The following figures (Figures 38–42) document completed results from techniques "
                "implemented in Sections 4, 5, and 7 of this report: knowledge distillation compression "
                "(Section 7.1), attentive feature selection (Section 4.4), self-supervised embeddings "
                "(Section 4.4), label-smoothed Mixup regularization (Section 4.2), and Opacus DP-SGD "
                "privacy training (Section 4.4). They are presented here for consolidated reference.")
if injected_386 in old_386:
    new_386 = old_386.replace(injected_386, "")
    if p386.runs:
        p386.runs[0].text = new_386
        for r in p386.runs[1:]: r.text = ''
    print("  Cleaned para 386 (active learning transitional sentence)")

# ─────────────────────────────────────────────────────────────────────────────
# STEP 3: Move Fig 41 (Mixup BCE) -> before para 233 "4.3 Hyperparameter Tuning"
# Add intro sentence + blank line, then figure block
print("\nMoving Fig 41 (Mixup BCE) -> before 4.3...")
intro_41 = new_para_elem(
    "Figure 41 below presents the training loss comparison between standard Binary Cross-Entropy "
    "(BCE) and Label-Smoothed Mixup loss, empirically demonstrating the regularization effect of "
    "convex input combination on our PyTorch wrapper models:"
)
anchor_4_3.addprevious(blank_para_elem())
for e in reversed([intro_41] + fig41_elems + [blank_para_elem()]):
    anchor_4_3.addprevious(e)
print("  Fig 41 moved.")

# ─────────────────────────────────────────────────────────────────────────────
# STEP 4: Move Fig 42 (Opacus DP-SGD) -> before "5.0 Insights and Interpretation"
# This places it at the end of Section 4, right before the insights section
print("Moving Fig 42 (Opacus DP-SGD) -> end of Section 4...")
intro_42 = new_para_elem(
    "Figure 42 below tracks the Opacus DP-SGD privacy budget (epsilon) consumption alongside the "
    "training loss profile across epochs, confirming that differential privacy guarantees are enforced "
    "without destabilizing the training loop:"
)
for e in reversed([intro_42] + fig42_elems + [blank_para_elem()]):
    anchor_5_0.addprevious(e)
print("  Fig 42 moved.")

# ─────────────────────────────────────────────────────────────────────────────
# STEP 5: Move Figs 39+40 (AttentiveTabNet + SCARF) -> before "5.3 Demographic Parity"
# This places them at the end of Section 5.2 Explainability
print("Moving Figs 39+40 (TabNet + SCARF) -> end of Section 5.2...")
intro_39 = new_para_elem(
    "Complementing the SHAP attributions above, Figure 39 presents the AttentiveTabNet instance-wise "
    "feature selection mask heatmap, and Figure 40 projects the SCARF self-supervised contrastive "
    "embeddings via t-SNE dimensionality reduction — both providing an additional lens on feature "
    "representational quality:"
)
for e in reversed([intro_39] + fig39_elems + fig40_elems + [blank_para_elem()]):
    anchor_5_3.addprevious(e)
print("  Figs 39+40 moved.")

# ─────────────────────────────────────────────────────────────────────────────
# STEP 6: Move Fig 38 (Knowledge Distillation) -> before "7.3 Summary of Excluded"
# This places it at the end of Section 7.2 Detailed Technical Specifications
print("Moving Fig 38 (Knowledge Distillation) -> end of Section 7.2...")
intro_38 = new_para_elem(
    "Figure 38 below presents the training loss curves comparing the teacher ensemble model against "
    "the compressed StudentNet, confirming successful knowledge transfer and stable convergence to "
    "the same loss plateau despite the dramatic size reduction:"
)
for e in reversed([intro_38] + fig38_elems + [blank_para_elem()]):
    anchor_7_3.addprevious(e)
print("  Fig 38 moved.")

# ─────────────────────────────────────────────────────────────────────────────
# STEP 7: Remove the fake "8.3 Supplementary Results" heading
print("\nRemoving fake '8.3 Supplementary' heading...")
heading_elem.getparent().remove(heading_elem)
print("  Removed.")

# ─────────────────────────────────────────────────────────────────────────────
# STEP 8: Save
print("\nSaving document...")
doc.save(DOCX_PATH)
print("Document saved successfully.")

# ─────────────────────────────────────────────────────────────────────────────
# STEP 9: Verify — reload and print relevant sections
print("\n=== VERIFICATION ===")
doc2 = Document(DOCX_PATH)

def has_image2(para):
    xml = etree.tostring(para._element, encoding='unicode')
    return '<w:drawing' in xml or '<v:shape' in xml

print("\n-- Around para 233 area (4.2 calibration -> 4.3 tuning) --")
count = 0
for i, p in enumerate(doc2.paragraphs):
    t = p.text.strip()
    img = '[IMG]' if has_image2(p) else ''
    if any(kw in p.text for kw in ["4.3", "4.2", "Brier", "Mixup", "Label-Smoothed", "Figure 41", "Isotonic", "Reliability", "Calibrat"]):
        print(f"  {i:4d} {img:5} {t[:110]}")
        count += 1
    if count > 20: break

print("\n-- Section 5.2 explainability area --")
count = 0
for i, p in enumerate(doc2.paragraphs):
    if any(kw in p.text for kw in ["5.2", "5.3", "AttentiveTabNet", "SCARF", "Figure 39", "Figure 40", "Demographic Parity"]):
        img = '[IMG]' if has_image2(p) else ''
        print(f"  {i:4d} {img:5} {p.text.strip()[:110]}")
        count += 1
    if count > 25: break

print("\n-- Section 7.2 end and 7.3 area --")
count = 0
for i, p in enumerate(doc2.paragraphs):
    if any(kw in p.text for kw in ["7.2", "7.3", "Knowledge Distill", "Figure 38", "StudentNet", "Summary of Excluded", "Implemented"]):
        img = '[IMG]' if has_image2(p) else ''
        print(f"  {i:4d} {img:5} {p.text.strip()[:110]}")
        count += 1
    if count > 20: break

print("\n-- Section 8.2/8.3 area (should be clean now) --")
for i, p in enumerate(doc2.paragraphs):
    if any(kw in p.text for kw in ["8.2", "8.3", "Figure 38", "Figure 39", "Figure 40", "Figure 41", "Figure 42", "Supplementary"]):
        img = '[IMG]' if has_image2(p) else ''
        print(f"  {i:4d} {img:5} {p.text.strip()[:110]}")
