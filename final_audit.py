import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx')

print('=== FINAL CLEAN REPORT AUDIT SUMMARY ===')
print(f'Total paragraphs: {len(doc.paragraphs)}')
print(f'Total tables: {len(doc.tables)}')

checks = [
    ('LightGBM as champion', lambda p: 'lightgbm (tuned) is designated' in p.lower()),
    ('Dynamic Champion Model', lambda p: 'dynamic champion model' in p.lower()),
    ('RobustScaler in preprocessing list', lambda p: 'normalization: applied a robustscaler' in p.lower()),
    ('116 features in preprocessing', lambda p: '116 input features' in p.lower() or '116 features' in p.lower()),
    ('Figure 18 correct champion', lambda p: 'dynamic champion model (lightgbm' in p.lower()),
    ('Simulator correct champion', lambda p: 'dynamic champion model in real-time' in p.lower()),
    ('ATE = 0.0', lambda p: 'ate is exactly 0.0' in p.lower() or 'average treatment effect (ate) is exactly 0.0' in p.lower()),
    ('14 surgical fixes', lambda p: '14 surgical fixes' in p.lower()),
    ('DML K-Fold cross-fitting', lambda p: 'k-fold cross-fitting' in p.lower()),
    ('Brier score 0.2426 to 0.2393', lambda p: '0.2426 to 0.2393' in p),
    ('SMOTE 24,120 balanced', lambda p: '24,120' in p),
]

for name, fn in checks:
    found = any(fn(p.text) for p in doc.paragraphs)
    status = 'FOUND' if found else 'MISSING'
    icon = 'OK' if found else 'WARN'
    print(f'  [{icon}] {name}: {status}')

print()
print('--- PROBLEMATIC TERMS ---')
problems = [
    ('StandardScaler in preprocessing (not comparison)',
     lambda p: 'normalization: applied a standardscaler' in p.lower()),
    ('selected best model (Random Forest)',
     lambda p: 'selected best model (random forest)' in p.lower()),
    ('XGBoost model in simulator',
     lambda p: 'evaluates the trained xgboost model' in p.lower()),
    ('113 columns in conclusion',
     lambda p: '113 columns' in p.lower()),
]
for name, fn in problems:
    found_paras = [i for i, p in enumerate(doc.paragraphs) if fn(p.text)]
    status = 'ISSUE' if found_paras else 'CLEAN'
    extra = str(found_paras) if found_paras else ''
    print(f'  [{status}] {name} {extra}')
