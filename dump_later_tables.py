import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'reports\WIA1006 Machine Learning - Tying the Data Knot Assignment Report V8 (final).docx')

print(f"Total tables in document: {len(doc.tables)}")
for i, table in enumerate(doc.tables):
    if i >= 10:
        print(f"=== Table {i} ===")
        for r_idx, row in enumerate(table.rows):
            row_txt = [cell.text.strip() for cell in row.cells]
            # Print unique text in row
            seen = set()
            unique_row = []
            for t in row_txt:
                if t not in seen:
                    seen.add(t)
                    unique_row.append(t)
            print(f"  Row {r_idx}: {unique_row}")
        print("-" * 50)
